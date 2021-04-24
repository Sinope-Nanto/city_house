import openpyxl
from openpyxl.styles import Side, Border,PatternFill, colors, Font, Alignment
from openpyxl.drawing.image import Image
import time

from docx.oxml.ns import qn
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt,Cm

class GenExcelReport:
    def __init__(self,url):
        self.edge = Side(style='thin', color='000000')
        self.border = Border(top=self.edge, bottom=self.edge, left=self.edge, right=self.edge)
        self.url = url
        self.report = openpyxl.Workbook()
    def create_report(self):
        indexsummary = self.report.active
        indexsummary.title = '指数汇总'
        totalradioplot = self.report.create_sheet()
        totalradioplot.title = '全国指数同比环比图'
        plotsimple = self.report.create_sheet()
        plotsimple.title = '作图（简版）'
        plotcomplex = self.report.create_sheet()
        plotcomplex.title = '作图（详细）'
        citysummary = self.report.create_sheet()
        citysummary.title = '城市汇总'
        citysort = self.report.create_sheet()
        citysort.title = '城市排序'
        cityindex = self.report.create_sheet()
        cityindex.title = '各城市指数'
        cityvolumn = self.report.create_sheet()
        cityvolumn.title = '各城市样本量'
        cityplot = self.report.create_sheet()
        cityplot.title = '各城市作图'
        return 
    def EndReport(self):
        self.report.save(self.url)

    def to_date(self, temp):
        year = int(temp)//12 + 6
        month = int(temp)%12 + 1
        return str(year).zfill(2) + '-' + str(month).zfill(2)
    
    def IndexSummary(self,**kwargs):
        # 参数列表:['volumn','east_volumn','mid_volumn','west_volumn',
        # 'under_90_volumn','90_144_volumn','above_144_volumn','volumn_year_on_year','volumn_chain',
        # 'index','east_index','mid_index','west_index',
        # 'year_on_year','chain','east_year_on_year','east_chain',
        # 'mid_year_on_year','mid_chain','west_year_on_year','west_chain',
        # 'index_under_90','index_90_144','index_above_144','year_on_year_under_90',
        # 'chain_under_90','year_on_year_90_144','chain_90_144','year_on_year_above_144','chain_above_144','index_CSJ']
        
        # 表格基本格式设置
        indexsummary = self.report['指数汇总']
        indexsummary.freeze_panes = 'C1'
        indexsummary.column_dimensions['A'].width = 20.0
        titlefont = Font(u'宋体',size = 12,bold=True,color='000000')
        datafont = Font(u'宋体',size = 10,bold=True,color='000000')
        grayfill = PatternFill("solid", fgColor='00C0C0C0')
        yellowfill = PatternFill("solid", fgColor='FFFFFF00')
        col_max = len(kwargs['volumn'])+2
        
        # 表头文字设置
        titlelist = ['交易量','指数汇总结果','区域性指数']
        rowlist = [1,15,43]
        for i in range(0,3):
            c = indexsummary.cell(row=rowlist[i],column=1)
            c.value = titlelist[i]
            c.font = titlefont
        
        # 表格颜色填充
        grayrow = [1,15,43]
        yellowrow = [24,25,28,29,33,38,39]
        for j in range(1,col_max):
            for i in grayrow:
                c = indexsummary.cell(row=i,column=j)
                c.fill = grayfill
            for i in yellowrow:
                c = indexsummary.cell(row=i,column=j)
                c.fill = yellowfill
        
        # 表格边框加粗
        borderrow = [i for i in range(2,10)] + [11,12] + [i for i in range(16,21)] + [i for i in range(22,30)]+ [31,32,33,34]+ [i for i in range(36,42)]+ [44,45]
        for i in borderrow:
            for j in range(1,col_max):
                c = indexsummary.cell(row=i,column=j)
                c.border = self.border

        # 日期填入
        datarow = [2,16,31,44]
        for j in datarow:
            for i in range(2,col_max):
                c = indexsummary.cell(row=j,column=i)
                c.value = self.to_date(i-2)
                c.font = datafont

        # 表头单元格合并
        indexsummary.merge_cells('A1:B1')
        indexsummary.merge_cells('A15:B15')
        indexsummary.merge_cells('A43:B43')

        # 填入表行名
        int_row_name = ['全国','东部','中部','西部','90以下','90-144','144以上']
        float_row_name = ['全国','东部','中部','西部','90以下','90-144','144以上','长三角']
        year_row_name = ['全国交易量同比变化','全国同比涨幅','东部同比涨幅','中部同比涨幅','西部同比涨幅',
        '90以下同比涨幅','90-144同比涨幅','144以上同比涨幅']
        chain_row_name = ['全国交易量环比变化','全国环比涨幅','东部环比涨幅','中部环比涨幅','西部环比涨幅',
        '90以下环比涨幅','90-144环比涨幅','144以上环比涨幅']
        int_row_list = [i for i in range(3,10)]
        float_row_list = [17,18,19,20,32,33,34,45]
        year_row_list = [11,22,24,26,28,36,38,40]
        chain_row_list = [12,23,25,27,29,37,39,41]

        rowname = int_row_name + float_row_name + year_row_name + chain_row_name
        rowlist = int_row_list + float_row_list + year_row_list + chain_row_list

        for i in range(0,len(rowname)):
            c = indexsummary.cell(row=rowlist[i],column=1)
            c.value = rowname[i]
            c.font = datafont


        # 填入整数型数据
        data_name = ['volumn','east_volumn','mid_volumn','west_volumn',
        'under_90_volumn','90_144_volumn','above_144_volumn']
        for i in range(0,len(int_row_list)):
            for j in range(2,col_max):
                c = indexsummary.cell(row=int_row_list[i],column=j)
                c.value = kwargs[data_name[i]][j-2]
                c.font = datafont

        # 填入浮点数型数据
        data_name = ['index','east_index','mid_index','west_index',
        'index_under_90','index_90_144','index_above_144','index_CSJ']
        for i in range(0,len(float_row_list)):
            for j in range(2,col_max):
                c = indexsummary.cell(row=float_row_list[i],column=j)
                c.value = float('%.2f'%kwargs[data_name[i]][j-2])
                c.font = datafont
        
        # 填入同比数据
        data_name = ['volumn_year_on_year',
        'year_on_year','east_year_on_year','mid_year_on_year','west_year_on_year',
        'year_on_year_under_90','year_on_year_90_144','year_on_year_above_144']
        for i in range(0,len(year_row_list)):
            for j in range(2,14):
                c = indexsummary.cell(row=year_row_list[i],column=j)
                c.value = '——'
                c.font = datafont
            for j in range(14,col_max):
                c = indexsummary.cell(row=year_row_list[i],column=j)
                c.value = '{:.2%}'.format(kwargs[data_name[i]][j-14])
                c.font = datafont
        
        # 填入环比数据
        data_name = ['volumn_chain',
        'chain','east_chain','mid_chain','west_chain',
        'chain_under_90','chain_90_144','chain_above_144']
        for i in range(0,len(year_row_list)):
            c = indexsummary.cell(row=chain_row_list[i],column=2)
            c.value = '——'
            c.font = datafont
            for j in range(3,col_max):
                c = indexsummary.cell(row=chain_row_list[i],column=j)
                c.value = '{:.2%}'.format(kwargs[data_name[i]][j-3])
                c.font = datafont
    
    def RadioPlot(self,**kwargs):
        # 参数列表:['index','year_on_year','chain','year_on_year_plot','chain_plot']

        # 表格基本格式设置
        radioplot = self.report['全国指数同比环比图']
        titlefont = Font(u'宋体',size = 12,bold=True,color='000000')
        datafont = Font(u'宋体',size = 10,bold=False,color='000000')
        bluefill = PatternFill("solid", fgColor='FFC0D9D9')
        rowmax = len(kwargs['index']) + 2

        # 填充颜色
        for i in range(1,5):
            c = radioplot.cell(row=1,column=i)
            c.fill = bluefill
        
        # 表格边框加粗
        for i in range(1,rowmax):
            for j in range(1,5):
                c = radioplot.cell(row=i,column=j)
                c.border = self.border
        
        # 填入表头
        c = radioplot.cell(row=1,column=2)
        c.value = '指数'
        c.font = titlefont
        c = radioplot.cell(row=1,column=3)
        c.value = '环比'
        c.font = titlefont
        c = radioplot.cell(row=1,column=4)
        c.value = '同比'
        c.font = titlefont

        # 填入日期
        for i in range(2,rowmax):
            c = radioplot.cell(row=i,column=1)
            c.value = self.to_date(i-2)
            c.font = titlefont
        
        # 填入数据
        for i in range(2,rowmax):
            c = radioplot.cell(row=i,column=2)
            c.value = float('%.2f'%kwargs['index'][i-2])
            c.font = datafont

            c = radioplot.cell(row=i,column=3)
            if i < 3:
                c.value = '——'
            else:
                c.value = float('%.2f'%(kwargs['chain'][i-3]*100))
            c.font = datafont

            c = radioplot.cell(row=i,column=4)
            if i < 14:
                c.value = '——'
            else:
                c.value = float('%.2f'%(kwargs['year_on_year'][i-14]*100))
            c.font = datafont
        
        # 添加图片
        img = Image(kwargs['year_on_year_plot'])
        img.width, img.height=1200, 300
        radioplot.add_image(img, 'E8')

        img = Image(kwargs['chain_plot'])
        img.width, img.height=1200, 300
        radioplot.add_image(img, 'E27')
    
    def SimplyPlot(self,**kwargs):
        # 参数列表:['url_index_plot','url_block_plot','url_line_plot']
        simplyplot = self.report['作图（简版）']
        simplyplot.column_dimensions['A'].height = 4.0
        titlefont = Font(u'宋体',size = 14,bold=True,color='000000')
        grayfill = PatternFill("solid", fgColor='00C0C0C0')

        # 填充颜色
        for j in [1,29,57]:
            for i in range(1,16):
                c = simplyplot.cell(row=j,column=i)
                c.fill = grayfill
        
        # 添加标题
        c = simplyplot.cell(row=1,column=1)
        c.value = '全国（简版封面&最简版图1）'
        c.font = titlefont
        c = simplyplot.cell(row=29,column=1)
        c.value = '各地区（最简版图2）'
        c.font = titlefont
        c = simplyplot.cell(row=57,column=1)
        c.value = '各面积（最简版图3）'
        c.font = titlefont

        # 添加图片
        img = Image(kwargs['url_index_plot'])
        img.width, img.height=1000, 500
        simplyplot.add_image(img, 'A2')
        img = Image(kwargs['url_block_plot'])
        img.width, img.height=1000, 500
        simplyplot.add_image(img, 'A30')
        img = Image(kwargs['url_line_plot'])
        img.width, img.height=1000, 500
        simplyplot.add_image(img, 'A58')

        # 合并单元格
        simplyplot.merge_cells('A1:F1')
        simplyplot.merge_cells('A29:F29')
        simplyplot.merge_cells('A57:F57')
    def CityVolumn(self,citylist :list,**kwargs):
        # 参数列表citylist中的城市的数据

        # 表格基本格式设置
        cityvolumn = self.report['各城市样本量']
        cityvolumn.freeze_panes = 'D3'
        cityvolumn.column_dimensions['A'].width = 5.0
        datafont = Font(u'宋体',size = 10,bold=False,color='000000')
        col_max = len(kwargs[citylist[0]]) + 3
        row_max = len(citylist) + 2

        # 填入表头时间
        for i in range(3,col_max):
            c = cityvolumn.cell(row=1,column=i)
            c.value = self.to_date(i-3)
            c.font = datafont
        # 填入序号
        for i in range(2,row_max):
            c = cityvolumn.cell(row=i,column=1)
            c.value = i-1
            c.font = datafont
            c = cityvolumn.cell(row=i,column=2)
            c.value = citylist[i-2]
            c.font = datafont
        # 填入数据
        for i in range(2,row_max):
            for j in range(3,col_max):
                c = cityvolumn.cell(row=i,column=j)
                c.value = kwargs[citylist[i-2]][j-3]
                c.font = datafont
    
    def CityIndex(self,citylist :list,**kwargs):
        # 参数列表citylist中的城市的数据

        # 表格基本格式设置
        cityindex = self.report['各城市指数']
        cityindex.freeze_panes = 'D3'
        cityindex.column_dimensions['A'].width = 5.0
        datafont = Font(u'宋体',size = 10,bold=False,color='000000')
        col_max = len(kwargs[citylist[0]]) + 3
        row_max = len(citylist) + 2

        # 填入表头时间
        for i in range(3,col_max):
            c = cityindex.cell(row=1,column=i)
            c.value = self.to_date(i-3)
            c.font = datafont
        # 填入序号
        for i in range(2,row_max):
            c = cityindex.cell(row=i,column=1)
            c.value = i-1
            c.font = datafont
            c = cityindex.cell(row=i,column=2)
            c.value = citylist[i-2]
            c.font = datafont
        # 填入数据
        for i in range(2,row_max):
            for j in range(3,col_max):
                c = cityindex.cell(row=i,column=j)
                c.value = float('%.2f'%kwargs[citylist[i-2]][j-3])
                c.font = datafont
    
    def CityPlot(self,citylist:list,poltlist:list):
        cityindex = self.report['各城市作图']
        datafont = Font(u'宋体',size = 10,bold=False,color='000000')
        
        # 25
        for i in range(0,len(citylist)):
            c = cityindex.cell(row=25*i + 1,column=1)
            c.value = '图' + str(i+1) + ' ' + citylist[i] + '\"城房指数\"'
            c.font = datafont
            img = Image(poltlist[i])
            img.width, img.height=1600, 400
            cityindex.add_image(img, 'A'+str(25*i+2))
    
    def CitySummary(self,informationlist:list):
        # list每项的参数列表:{'city_name','index_this_month','index_last_month','index_last_year',
        # 'chain_radio','year_on_year','volumn_chain','volumn_year'}
        
        # 表格基本格式设置
        citysummary = self.report['城市汇总']
        citysummary.freeze_panes = 'C2'
        citysummary.column_dimensions['A'].width = 5.0
        datafont = Font(u'宋体',size = 10,bold=False,color='000000')
        col_max = 10
        row_max = len(informationlist) + 2

        # 表格边框加粗
        for i in range(1,row_max):
            for j in range(1,col_max):
                c = citysummary.cell(row=i,column=j)
                c.border = self.border
        
        # 填入表头
        col_name_list = ['当月指数值','上月指数值','去年同月指数值',
        '环比','同比','交易量环比','交易量同比']
        for i in range(3,col_max):
            c = citysummary.cell(row=1,column=i)
            c.value = col_name_list[i-3]
            c.font = datafont
        
        # 填入数据
        col_data_list = ['city_name','index_this_month','index_last_month','index_last_year',
        'chain_radio','year_on_year','volumn_chain','volumn_year']
        for i in range(2,row_max):
            c = citysummary.cell(row=i,column=1)
            c.value = i-1
            c.font = datafont
            for j in range(2,col_max):
                c = citysummary.cell(row=i,column=j)
                if j == 2:
                    c.value = informationlist[i-2]['city_name']
                elif j < 6:
                    c.value = float('%.2f'%informationlist[i-2][col_data_list[j-2]])
                else:
                    c.value = '%.2f'%(informationlist[i-2][col_data_list[j-2]]*100) + '%'
                c.font = datafont

    def ComplexPlot(self,**kwargs):
        # 参数列表:['url_index_plot','url_block_plot','url_line_plot']
        complexplot = self.report['作图（详细）']

        # 添加图片
        img = Image(kwargs['url_index_plot'])
        img.width, img.height=1600, 400
        complexplot.add_image(img, 'A1')
        img = Image(kwargs['url_block_plot'])
        img.width, img.height=1600, 400
        complexplot.add_image(img, 'A26')
        img = Image(kwargs['url_line_plot'])
        img.width, img.height=1600, 400
        complexplot.add_image(img, 'A51')
    
    def CitySort(self,month:int,cityinformation:list):
        # 列表中每个元素应有参数:{'city_name','chain0','chain1' ,'chain2',
        # 'chain3' 'year0' ,'year1' ,'year2','year3'}

        # 表格基本格式设置
        citysort = self.report['城市排序']
        datafont = Font(u'宋体',size = 10,bold=True,color='000000')
        Reddatafont = Font(u'宋体',size = 10,bold=False,color='FF0000')
        bluefill = PatternFill("solid", fgColor='FF3299CC')
        rowmax = len(cityinformation) + 3
        colmax = 11
        align = Alignment(horizontal='center', vertical='center')

        # 表格边框加粗
        for i in range(1,rowmax):
            for j in range(1,colmax):
                c = citysort.cell(row=i,column=j)
                c.border = self.border

        # 颜色填充
        for j in range(1,colmax):
            for i in [1,2]:
                c = citysort.cell(row=i,column=j)
                c.fill = bluefill
        
        # 添加表头
        c = citysort.cell(row=1,column=3)
        c.value = '环比'
        c.font = datafont
        c.alignment = align
        citysort.merge_cells('C1:F1')

        c = citysort.cell(row=1,column=7)
        c.value = '同比'
        c.font = datafont
        c.alignment = align
        citysort.merge_cells('G1:J1')

        c = citysort.cell(row=2,column=1)
        c.value = '编号'
        c.font = datafont

        c = citysort.cell(row=2,column=2)
        c.value = '城市'
        c.font = datafont

        for i in range(0,4):
            c = citysort.cell(row=2,column=3+i)
            if month-i > 0:
                c.value = str(month-i) + '月'
            else:
                c.value = str(12+month-i) + '月'
            c.font = datafont
            c = citysort.cell(row=2,column=7+i)
            if month-i > 0:
                c.value = str(month-i) + '月'
            else:
                c.value = str(12+month-i) + '月'
            c.font = datafont
        
        # 填入数据
        col_data_list = ['city_name','chain0','chain1' ,'chain2','chain3' ,'year0' ,'year1' ,'year2','year3']
        for i in range(3,rowmax):
            c = citysort.cell(row=i,column=1)
            c.value = i-2
            c.font = datafont
            for j in range(2,colmax):
                c = citysort.cell(row=i,column=j)
                if j == 2:
                    c.value = cityinformation[i-3]['city_name']
                else:
                    c.value = float('%.2f'%(cityinformation[i-3][col_data_list[j-2]]*100))
                    if c.value == 0 or c.value == -100:
                        c.value = ''
                    elif c.value < 0:
                        c.font = Reddatafont

        
        # 添加排序功能
        citysort.auto_filter.ref = "A2:J" + str(colmax-1)
        citysort.auto_filter.add_sort_condition("A3:A"+ str(colmax -1))

class GenWordReport:
    def __init__(self,url):
        self.url = url
        self.doc = Document()
    
    #生成封面
    def Firstpage(self, year:int, month:int, **kwargs):
        title = str(year) + '年' + str(month) + '月全国40个重点城市“城房指数”月报'
        self.doc.add_paragraph('')
        self.doc.add_paragraph('')
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run(title)
        tit.font.name = '黑体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        tit.font.size = Pt(28)
        tit.font.bold = True

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('（试行）')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(20)
        self.doc.add_paragraph('')

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('全国40个重点城市各月“城房指数”汇总值')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        p = self.doc.add_paragraph()
        pic = p.add_run().add_picture(kwargs['index_image_url'])
        pic.height = Cm(10.96)
        pic.width =  Cm(17.65)

        self.doc.add_section(start_type = None)
    
    def Secondpage(self, year:int, month:int, citylist:list, **kwargs):
        #参数列表:['index', 'chain', 'year_on_year', 'city_name', 'city_chain', 'city_year_on_year']
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('简要说明')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(20)
        tit.font.bold = True

        p = self.doc.add_paragraph()
        l = len(kwargs['index'])
        index = float('%.2f'%kwargs['index'][l - 1])
        chain = float('%.2f'%(kwargs['chain'][l - 1]*100))
        year_on_year = float('%.2f'%(kwargs['year_on_year'][l - 1]*100))
        s = 0
        if chain < 0:
            str1 = '下降' + str(-chain) + '%'
        else:
            str1 = '上涨' + str(chain) + '%'
        if year_on_year < 0:
            str2 = '下降' + str(-year_on_year) + '%。'
        else:
            str2 = '上涨' + str(year_on_year) + '%。'
        zero = []
        for i in range(len(citylist)):
            if float(kwargs['city_index'][i]) == 0:
                zero.append(i)
        if len(zero) == 0:
            s = 0
            add = p.add_run('  ' + str(year) + '年' + str(month) + '月”城房指数“全国汇总值为' + str(index) + '点，同比' + str1 + '，环比' + str2)
            add.font.name = '仿宋_GB2312'
            add.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            add.font.size = Pt(14)
        else:
            s = 1
            add = p.add_run('  ' + str(year) + '年' + str(month) + '月”城房指数“全国')
            add.font.name = '仿宋_GB2312'
            add.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            add.font.size = Pt(14)

            add = p.add_run('1')
            add.font.superscript = True
            add.font.name = '仿宋_GB2312'
            add.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            add.font.size = Pt(14)

            add = p.add_run('汇总值为' + str(index) + '点，同比' + str1 + '，环比' + str2)
            add.font.name = '仿宋_GB2312'
            add.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            add.font.size = Pt(14)
        #统计环比涨跌

        l = len(citylist)
        higher = []
        num_higher = 0
        higher_than_5 = []
        num_lower = 0
        lower = []
        lower_than_5 = []
        for i in range(l):
            if float(kwargs['city_chain'][i]) > 0:
                higher.append(i)
                num_higher+=1
                if float(kwargs['city_chain'][i]) > 0.05:
                    higher_than_5.append(i)
            elif float(kwargs['city_chain'][i]) < 0:
                lower.append(i)
                num_lower+=1
                if float(kwargs['city_chain'][i]) < -0.05:
                    lower_than_5.append(i)
        #环比涨幅

        p = self.doc.add_paragraph()
        if num_higher == 0:
            tit = p.add_run('  没有任何一个城市的“城房指数”环比上涨；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('  有' + str(num_higher) + '个城市的“城房指数”环比上涨，')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

            if len(higher_than_5) == 0:
                tit = p.add_run('但涨幅均在5%以内；')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
            else:
                tit = p.add_run('其中')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)

                for i in range(len(higher_than_5)):
                    s += 1
                    str1 = '(' + str(float('%.2f' % (kwargs['city_chain'][higher_than_5[i]] * 100)))
                    if i != (len(higher_than_5) - 1):
                        tit = p.add_run(citylist[higher_than_5[i]])
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str(s))
                        tit.font.superscript = True
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str1 + '%）、')
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)
                    else:
                        tit = p.add_run(citylist[higher_than_5[i]])
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str(s))
                        tit.font.superscript = True
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str1 + '%）')
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                tit = p.add_run(str(len(higher_than_5)) + '个城市涨幅超过5%；')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
        #环比跌幅

        if num_lower == 0:
            tit = p.add_run('  没有任何一个城市的“城房指数”环比下降；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('有' + str(num_lower) + '个城市的“城房指数”环比下降，')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

            if len(lower_than_5) == 0:
                tit = p.add_run('但降幅均在5%以内；')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
            else:
                tit = p.add_run('其中')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)

                for i in range(len(lower_than_5)):
                    s += 1
                    str1 = '(' + str(float('%.2f' % (kwargs['city_chain'][lower_than_5[i]] * 100)))
                    if i != (len(lower_than_5) - 1):
                        tit = p.add_run(citylist[lower_than_5[i]])
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str(s))
                        tit.font.superscript = True
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str1 + '%）、')
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)
                    else:
                        tit = p.add_run(citylist[lower_than_5[i]])
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str(s))
                        tit.font.superscript = True
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str1 + '%）')
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                tit = p.add_run(str(len(lower_than_5)) + '个城市降幅超过5%。')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
        #统计同比涨降

        higher = []
        num_higher = 0
        higher_than_15 = []
        num_lower = 0
        lower = []
        lower_than_15 = []
        for i in range(l):
            if float(kwargs['city_year_on_year'][i]) > 0:
                higher.append(i)
                num_higher += 1
                if float(kwargs['city_year_on_year'][i]) > 0.15:
                    higher_than_15.append(i)
            elif float(kwargs['city_year_on_year'][i]) < 0:
                lower.append(i)
                num_lower += 1
                if float(kwargs['city_year_on_year'][i]) < -0.15:
                    lower_than_15.append(i)
        # 同比涨幅

        p = self.doc.add_paragraph()
        if num_higher == 0:
            tit = p.add_run('  没有任何一个城市的“城房指数”同比上涨；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('  有' + str(num_higher) + '个城市的“城房指数”同比上涨，')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

            if len(higher_than_15) == 0:
                tit = p.add_run('但涨幅均在15%以内；')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
            else:
                tit = p.add_run('其中')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)

                for i in range(len(higher_than_15)):
                    str1 = '(' + str(float('%.2f' % (kwargs['city_year_on_year'][higher_than_15[i]] * 100)))
                    if i != (len(higher_than_15) - 1):
                        tit = p.add_run(citylist[higher_than_15[i]] + str1 + '%）、')
                    else:
                        tit = p.add_run(citylist[higher_than_15[i]] + str1 + '%）')
                    tit.font.name = '仿宋_GB2312'
                    tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    tit.font.size = Pt(14)

                tit = p.add_run(str(len(higher_than_15)) + '个城市涨幅超过15%；')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
        # 同比跌幅

        if num_lower == 0:
            tit = p.add_run('  没有任何一个城市的“城房指数”同比下降；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('有' + str(num_lower) + '个城市的“城房指数”同比下降，')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

            if len(lower_than_15) == 0:
                tit = p.add_run('但降幅均在15%以内；')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
            else:
                tit = p.add_run('其中')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)

                for i in range(len(lower_than_15)):
                    str1 = '(' + str(float('%.2f' % (kwargs['city_year_on_year'][lower_than_15[i]] * 100)))
                    if i != (len(lower_than_15) - 1):
                        tit = p.add_run(citylist[lower_than_15[i]] + str1 + '%）、')
                    else:
                        tit = p.add_run(citylist[lower_than_15[i]] + str1 + '%）')
                    tit.font.name = '仿宋_GB2312'
                    tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    tit.font.size = Pt(14)

                tit = p.add_run(str(len(lower_than_15)) + '个城市降幅超过15%。')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
        #插入图片

        p = self.doc.add_paragraph()
        pic = p.add_run().add_picture(kwargs['chain_image_url'])
        pic.height = Cm(5.46)
        pic.width = Cm(14.67)

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('图1 40城市“城房指数”同比变化图')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True
        # 脚注

        p = self.doc.add_paragraph()
        tit = p.add_run('注：')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(9)
        tit.font.bold = True
        # 脚注正文

        if len(zero) != 0:
            p = self.doc.add_paragraph()
            tit = p.add_run('1')
            tit.font.name = '宋体'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            tit.font.size = Pt(9)

            for i in range(len(zero)):
                if i != len(zero) - 1:
                    tit = p.add_run(citylist[zero[i]] + '、')
                    tit.font.name = '宋体'
                    tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    tit.font.size = Pt(9)
                else:
                    tit = p.add_run(citylist[zero[i]] + ':因指数软件问题无法计算，清华技术已经在研发新的软件来替换现有软件的不足。')
                    tit.font.name = '宋体'
                    tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    tit.font.size = Pt(9)

        self.doc.add_page_break()

        p = self.doc.add_paragraph()
        pic = p.add_run().add_picture(kwargs['yearonyear_image_url'])
        pic.height = Cm(5.23)
        pic.width = Cm(14.65)

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('图2 40城市“城房指数”环比变化图')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        self.doc.add_section(start_type = None)

    def Charts(self, year:int, month:int, citylist:list, **kwargs):

        #数据从2008年1月开始
        p = self.doc.add_paragraph()
        section = self.doc.sections[2]
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)
        if  int(((year - 2008) * 12 + month) // 2) == 0:
            row_num = int(((year - 2008) * 12 + month) // 2)
        else:
            row_num = int(((year - 2008) * 12 + month) // 2) + 1
        r_num = row_num
        row_num = int((int((row_num - 24) // 25) + 2) + int(row_num))
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('表1 全国40个重点城市“城房指数”各月汇总值')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        #表1
        table = self.doc.add_table(row_num, 8)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.width = Cm(18.62)
        data_row = 1
        for i in range(row_num):
            if  ((i == 0) or (((i + 1) % 26) == 0)):

                p = table.cell(i, 0).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'月份')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 1).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'指数值')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 2).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'环比')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 3).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'同比')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 4).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'月份')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 5).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'指数值')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 6).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'环比')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 7).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'同比')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True


            else:
                p = table.cell(i, 0).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                m = str(int(data_row) % 12)
                if m == '0':
                    m = '12'
                y = (int(data_row) // 12) + 2008
                if (int(data_row) % 12) == 0:
                    y -= 1
                y = str(y)
                str1 = y + '年' + m + '月'
                run = p.add_run(u'%s'%str1)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 1).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f'%kwargs['index'][data_row - 1]))
                run = p.add_run(u'%s'%data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 2).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % (float(kwargs['chain'][data_row - 1]) * 100))) + "%"
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 3).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % (float(kwargs['year_on_year'][data_row - 1]) * 100))) + "%"
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                if  i != row_num - 1:
                    p = table.cell(i, 4).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    m = str(int(data_row + r_num) % 12)
                    if m == '0':
                        m = '12'
                    y = (int(data_row + r_num) // 12) + 2008
                    if (int(data_row + r_num) % 12) == 0:
                        y -= 1
                    y = str(y)
                    str1 = y + '年' + m + '月'
                    run = p.add_run(u'%s' % str1)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                    p = table.cell(i, 5).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(float('%.2f' % kwargs['index'][data_row + r_num - 1]))
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                    p = table.cell(i, 6).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(float('%.2f' % (float(kwargs['chain'][data_row + r_num - 1]) * 100))) + "%"
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                    p = table.cell(i, 7).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(float('%.2f' % (float(kwargs['year_on_year'][data_row + r_num - 1]) * 100))) + "%"
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                data_row += 1

        self.doc.add_page_break()

        #表2
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('表2 %s年%s月40个重点城市“城房指数”'%(str(year),str(month)))
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        table = self.doc.add_table(21, 8)

        #kwargs['city_index']:当月城市“城房指数”值的列表，按城市顺序排列，'city_chain'以及'city_year_on_year'以此类推
        for i in range(21):
            if i != 0:
                p = table.cell(i, 0).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(citylist[i - 1])     
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 1).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f'%kwargs['city_index'][i - 1]))
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 2).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % (float(kwargs['city_chain'][i - 1]) * 100))) + "%"
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 3).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % (float(kwargs['city_year_on_year'][i - 1]) * 100))) + "%"
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 4).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(citylist[i + 19])
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 5).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % kwargs['city_index'][i + 19]))
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 6).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % (float(kwargs['city_chain'][i + 19]) * 100))) + "%"
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 7).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % (float(kwargs['city_year_on_year'][i + 19]) * 100))) + "%"
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

            else:
                p = table.cell(i, 0).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '城市'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 1).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '指数值'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 2).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '环比'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 3).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '同比'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 4).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '城市'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 5).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '指数值'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 6).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '环比'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 7).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '同比'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

        self.doc.add_page_break()

        #表3
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('表3 %s年%s月40个重点城市“城房指数”样本量' % (str(year), str(month)))
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        table = self.doc.add_table(21, 8)

        # kwargs['city_index']:当月城市“城房指数”值的列表，按城市顺序排列，'city_chain'以及'city_year_on_year'以此类推
        # kwargs['city_index_1']:上月城市“城房指数”值的列表，按城市顺序排列
        # kwargs['city_index_2']:两个月前城市“城房指数”值的列表，按城市顺序排列
        for i in range(21):
            if i != 0:
                p = table.cell(i, 0).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(citylist[i - 1])
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 1).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % kwargs['city_index'][i - 1]))
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 2).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % kwargs['city_index_1'][i - 1]))
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 3).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % kwargs['city_index_2'][i - 1]))
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 4).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(citylist[i + 19])
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 5).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % kwargs['city_index'][i + 19]))
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 6).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % kwargs['city_index_1'][i + 19]))
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 7).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % kwargs['city_index_2'][i + 19]))
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

            else:
                if int(month) == 2:
                    mon = [12, 1, 2]
                elif int(month) == 1:
                    mon = [11, 12, 1]
                else:
                    mon = [int(month) - 2, int(month) - 1, int(month)]

                p = table.cell(i, 0).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '城市'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 1).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(mon[0]) + '月样本量'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 2).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(mon[1]) + '月样本量'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 3).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(mon[2]) + '月样本量'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 4).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '城市'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 5).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(mon[0]) + '月样本量'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 6).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(mon[1]) + '月样本量'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 7).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(mon[2]) + '月样本量'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True
        self.doc.add_section(start_type = None)

    def Attach(self, year:int, month:int, citylist:list, **kwargs):

        section = self.doc.sections[3]
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('附：' + str(year) + '年' + str(month) + '月各子市场“城房指数”变化情况')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        #一
        p = self.doc.add_paragraph()
        tit = p.add_run('一、东中西区域子市场')
        tit.font.name = '黑体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        tit.font.size = Pt(14)

        p = self.doc.add_paragraph()
        tit = p.add_run('   东中西区域子市场中，东部地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(kwargs['east_index'][len(kwargs['east_index']) - 1]) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['east_year_on_year'][len(kwargs['east_index']) - 1]) > 0:
            tit = p.add_run('上涨' + str(float('%.2f' %(kwargs['east_year_on_year'][len(kwargs['east_index']) - 1])) * 100) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-kwargs['east_year_on_year'][len(kwargs['east_index']) - 1])) * 100) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['east_chain'][len(kwargs['east_index']) - 1]) > 0:
            tit = p.add_run('上涨' + str(float('%.2f' %(kwargs['east_chain'][len(kwargs['east_index']) - 1])) * 100) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-kwargs['east_chain'][len(kwargs['east_index']) - 1])) * 100) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        tit = p.add_run('中部地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(kwargs['mid_index'][len(kwargs['mid_index']) - 1]) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['mid_year_on_year'][len(kwargs['mid_index']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (kwargs['mid_year_on_year'][len(kwargs['mid_index']) - 1])) * 100) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-kwargs['mid_year_on_year'][len(kwargs['mid_index']) - 1])) * 100) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['mid_chain'][len(kwargs['mid_index']) - 1]) > 0:
            tit = p.add_run('上涨' + str(float('%.2f' % (kwargs['mid_chain'][len(kwargs['mid_index']) - 1])) * 100) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-kwargs['mid_chain'][len(kwargs['mid_index']) - 1])) * 100) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        tit = p.add_run('西部地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(kwargs['west_index'][len(kwargs['west_index']) - 1]) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['west_year_on_year'][len(kwargs['west_index']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (kwargs['west_year_on_year'][len(kwargs['west_index']) - 1])) * 100) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-kwargs['west_year_on_year'][len(kwargs['west_index']) - 1])) * 100) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['west_chain'][len(kwargs['west_index']) - 1]) > 0:
            tit = p.add_run('上涨' + str(float('%.2f' % (kwargs['west_chain'][len(kwargs['west_index']) - 1])) * 100) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-kwargs['west_chain'][len(kwargs['west_index']) - 1])) * 100) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        p = self.doc.add_paragraph()
        pic = p.add_run().add_picture(kwargs['index_by_block_image_url'])
        pic.height = Cm(6.17)
        pic.width = Cm(14.63)

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('附图1 东中西区域子市场走势图')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        #二

        p = self.doc.add_paragraph()
        tit = p.add_run('二、各面积子市场')
        tit.font.name = '黑体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        tit.font.size = Pt(14)

        p = self.doc.add_paragraph()
        tit = p.add_run('   各面积子市场中，“90平方米以下”指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(kwargs['index_under_90'][len(kwargs['index_under_90']) - 1]) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['year_on_year_under_90'][len(kwargs['index_under_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (kwargs['year_on_year_under_90'][len(kwargs['index_under_90']) - 1])) * 100) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-kwargs['year_on_year_under_90'][len(kwargs['index_under_90']) - 1])) * 100) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chain_under_90'][len(kwargs['index_under_90']) - 1]) > 0:
            tit = p.add_run('上涨' + str(float('%.2f' % (kwargs['chain_under_90'][len(kwargs['index_under_90']) - 1])) * 100) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-kwargs['chain_under_90'][len(kwargs['index_under_90']) - 1])) * 100) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        tit = p.add_run('“90至144平方米”指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(kwargs['index_90_144'][len(kwargs['index_90_144']) - 1]) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['year_on_year_90_144'][len(kwargs['index_90_144']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (kwargs['year_on_year_90_144'][len(kwargs['index_90_144']) - 1])) * 100) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-kwargs['year_on_year_90_144'][len(kwargs['index_90_144']) - 1])) * 100) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chain_90_144'][len(kwargs['index_90_144']) - 1]) > 0:
            tit = p.add_run('上涨' + str(float('%.2f' % (kwargs['chain_90_144'][len(kwargs['index_90_144']) - 1])) * 100) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-kwargs['chain_90_144'][len(kwargs['index_90_144']) - 1])) * 100) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        tit = p.add_run('“144平方米以上”指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(kwargs['index_above_144'][len(kwargs['index_above_144']) - 1]) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['year_on_year_above_144'][len(kwargs['index_above_144']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (kwargs['year_on_year_above_144'][len(kwargs['index_above_144']) - 1])) * 100) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-kwargs['year_on_year_above_144'][len(kwargs['index_above_144']) - 1])) * 100) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chain_above_144'][len(kwargs['index_above_144']) - 1]) > 0:
            tit = p.add_run('上涨' + str(float('%.2f' % (kwargs['chain_above_144'][len(kwargs['index_above_144']) - 1])) * 100) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-kwargs['chain_above_144'][len(kwargs['index_above_144']) - 1])) * 100) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        p = self.doc.add_paragraph()
        pic = p.add_run().add_picture(kwargs['index_by_buildarea_image_url'])
        pic.height = Cm(6.17)
        pic.width = Cm(14.63)

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('附图2 各面积子市场走势图')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True
        #附注

        p = self.doc.add_paragraph()
        tit = p.add_run('附注')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(10)
        tit.font.bold = True

        p = self.doc.add_paragraph()
        tit = p.add_run('东部地区包括北京、天津、石家庄、沈阳、大连、上海、南京、无锡、苏州、杭州、宁波、温州、福州、厦门、济南、青岛、广州、深圳、海'
                        '口、三亚等20个城市。中部地区包括太原、长春、哈尔滨、合肥、南昌、郑州、武汉、长沙等8个城市。西部地区包括呼和浩特、南宁、北海'
                        '、重庆、成都、贵阳、昆明、西安、兰州、西宁、银川、乌鲁木齐等12个城市。')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(10)
    
    def EndReport(self):
        self.doc.save(self.url)