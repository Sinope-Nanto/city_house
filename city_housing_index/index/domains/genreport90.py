import openpyxl
from openpyxl.styles import Side, Border, PatternFill, colors, Font, Alignment
from openpyxl.drawing.image import Image
import time

from docx.oxml.ns import qn
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm


class GenExcelReport90:
    def __init__(self, url):
        self.edge = Side(style='thin', color='000000')
        self.border = Border(top=self.edge, bottom=self.edge, left=self.edge, right=self.edge)
        self.url = url
        self.report = openpyxl.Workbook()

    def create_report(self):
        indexsummary = self.report.active
        indexsummary.title = '指数汇总'
        totalradioplot = self.report.create_sheet()
        totalradioplot.title = '全国指数同比环比图'
        plotsimple = self.report.create_sheet()
        plotsimple.title = '作图（简版）'
        plotcomplex = self.report.create_sheet()
        plotcomplex.title = '作图（详细）'
        citysummary = self.report.create_sheet()
        citysummary.title = '城市汇总'
        cityline = self.report.create_sheet()
        cityline.title = '各线城市汇总'
        citysort = self.report.create_sheet()
        citysort.title = '城市排序'
        cityindex = self.report.create_sheet()
        cityindex.title = '各城市指数'
        cityvolumn = self.report.create_sheet()
        cityvolumn.title = '各城市样本量'
        cityplot = self.report.create_sheet()
        cityplot.title = '各城市作图'
        line = self.report.create_sheet()
        line.title = '各线城市子市场'

        return

    def EndReport(self):
        self.report.save(self.url)

    def to_date(self, temp):
        year = int(temp) // 12 + 9
        month = int(temp) % 12 + 1
        return str(year).zfill(2) + '-' + str(month).zfill(2)

    def IndexSummary(self, **kwargs):
        # 参数列表:['volumn', 'east_volumn', 'mid_volumn', 'west_volumn',
        #             'under_90_volumn', '90_144_volumn', 'above_144_volumn',
        #             'DONGBEI_volumn', 'HUABEI_volumn', 'HUADONG_volumn', 'HUAZHONG_volumn', 'HUANAN_volumn', 'XINAN_volumn', 'XIBEI_volumn', 
        #             'firstline_volumn', 'secondline_volumn', 'thirdline_volumn', 'forth_volumn', 'ZSJ_volumn', 'CSJ_volumn', 'HBH_volumn',
        #             'index', 'east_index', 'mid_index', 'west_index',
        #             'index_under_90', 'index_90_144', 'index_above_144',
        #             'DONGBEI_index', 'HUABEI_index', 'HUADONG_index', 'HUAZHONG_index', 'HUANAN_index', 'XINAN_index', 'XIBEI_index', 
        #             'firstline_index', 'secondline_index', 'thirdline_index', 'forth_index', 'ZSJ_index', 'CSJ_index', 'HBH_index',
        #             'volumn_year_on_year',
        #              'year_on_year', 'east_year_on_year', 'mid_year_on_year', 'west_year_on_year',
        #              'year_on_year_under_90', 'year_on_year_90_144', 'year_on_year_above_144',
        #              'DONGBEI_year_on_year', 'HUABEI_year_on_year', 'HUADONG_year_on_year', 'HUAZHONG_year_on_year', 'HUANAN_year_on_year', 'XINAN_year_on_year', 'XIBEI_year_on_year', 
        #              'firstline_year_on_year', 'secondline_year_on_year', 'thirdline_year_on_year', 'forth_year_on_year', 'ZSJ_year_on_year', 'CSJ_year_on_year', 'HBH_year_on_year',
        #              'volumn_chain',
        #              'chain', 'east_chain', 'mid_chain', 'west_chain',
        #              'chain_under_90', 'chain_90_144', 'chain_above_144',
        #              'DONGBEI_chain', 'HUABEI_chain', 'HUADONG_chain', 'HUAZHONG_chain', 'HUANAN_chain', 'XINAN_chain', 'XIBEI_chain', 
        #              'firstline_chain', 'secondline_chain', 'thirdline_chain', 'forth_chain', 'ZSJ_chain', 'CSJ_chain', 'HBH_chain']

        # 表格基本格式设置
        indexsummary = self.report['指数汇总']
        indexsummary.freeze_panes = 'C1'
        indexsummary.column_dimensions['A'].width = 20.0
        titlefont = Font(u'宋体', size=12, bold=True, color='000000')
        datafont = Font(u'宋体', size=10, bold=False, color='000000')
        grayfill = PatternFill("solid", fgColor='00C0C0C0')
        yellowfill = PatternFill("solid", fgColor='FFFFFF00')
        col_max = len(kwargs['volumn']) + 2

        # 表头文字设置
        titlelist = ['交易量', '指数汇总结果', '新子指数系统交易量','新子指数系统指数结果']
        rowlist = [1, 15, 43, 60]
        for i in range(0, 4):
            c = indexsummary.cell(row=rowlist[i], column=1)
            c.value = titlelist[i]
            c.font = titlefont

        # 表格颜色填充
        grayrow = [1, 15, 43, 60]
        yellowrow = [24, 25, 28, 29, 33, 38, 39, 70, 71, 74, 75, 78, 79, 
        82, 83, 91, 92, 95, 96, 105, 106, 109, 110]
        for j in range(1, col_max):
            for i in grayrow:
                c = indexsummary.cell(row=i, column=j)
                c.fill = grayfill
            for i in yellowrow:
                c = indexsummary.cell(row=i, column=j)
                c.fill = yellowfill

        # 表格边框加粗
        borderrow = ([i for i in range(2, 10)] + [11, 12] + [i for i in range(16, 21)] + [i for i in range(22, 30)]
        + [31, 32, 33, 34] + [i for i in range(36, 42)] + [i for i in range(44, 59)] + [i for i in range(61, 69)] 
        + [i for i in range(70, 84)] + [i for i in range(85, 90)] + [i for i in range(91, 99)] + [i for i in range(100, 104)]
        + [i for i in range(105, 111)])
        for i in borderrow:
            for j in range(1, col_max):
                c = indexsummary.cell(row=i, column=j)
                c.border = self.border

        # 日期填入
        datarow = [2, 16, 31, 44, 61, 85, 100]
        for j in datarow:
            for i in range(2, col_max):
                c = indexsummary.cell(row=j, column=i)
                c.value = self.to_date(i - 2)
                c.font = datafont

        # 表头单元格合并
        indexsummary.merge_cells('A1:B1')
        indexsummary.merge_cells('A15:B15')
        indexsummary.merge_cells('A43:B43')
        indexsummary.merge_cells('A60:B60')

        # 填入表行名
        int_row_name = ['全国', '东部', '中部', '西部', '90以下', '90-144', '144以上', '东北', '华北', '华东', '华中', '华南'
        , '西南', '西北', '一线城市', '二线城市', '三线城市', '四线城市', '珠三角', '长三角', '环渤海']
        float_row_name = ['全国', '东部', '中部', '西部', '90以下', '90-144', '144以上', '东北', '华北', '华东', '华中', '华南'
        , '西南', '西北', '一线城市', '二线城市', '三线城市', '四线城市', '珠三角', '长三角', '环渤海']
        year_row_name = ['全国交易量同比变化', '全国同比涨幅', '东部同比涨幅', '中部同比涨幅', '西部同比涨幅',
                         '90以下同比涨幅', '90-144同比涨幅', '144以上同比涨幅', '东北同比涨幅', '华北同比涨幅', '华东同比涨幅', '华中同比涨幅', '华南同比涨幅',
                         '西南同比涨幅', '西北同比涨幅', '一线城市同比涨幅', '二线城市同比涨幅', '三线城市同比涨幅', '四线城市同比涨幅', 
                         '珠三角同比涨幅', '长三角同比涨幅', '环渤海同比涨幅']
        chain_row_name = ['全国交易量环比变化', '全国环比涨幅', '东部环比涨幅', '中部环比涨幅', '西部环比涨幅',
                          '90以下环比涨幅', '90-144环比涨幅', '144以上环比涨幅', '东北环比涨幅', '华北环比涨幅', '华东环比涨幅', '华中环比涨幅', '华南环比涨幅',
                         '西南环比涨幅', '西北环比涨幅', '一线城市环比涨幅', '二线城市环比涨幅', '三线城市环比涨幅', '四线城市环比涨幅', 
                         '珠三角环比涨幅', '长三角环比涨幅', '环渤海环比涨幅']
        int_row_list = [i for i in range(3, 10)] + [i for i in range(45, 59)]
        float_row_list = [17, 18, 19, 20, 32, 33, 34] + [i for i in range(62, 69)] + [i for i in range(86, 90)] + [101, 102, 103] 
        year_row_list = [11, 22, 24, 26, 28, 36, 38, 40, 70, 72, 74, 76, 78, 80, 82, 91, 93, 95, 97, 105, 107, 109]
        chain_row_list = [(i + 1) for i in year_row_list]

        rowname = int_row_name + float_row_name + year_row_name + chain_row_name
        rowlist = int_row_list + float_row_list + year_row_list + chain_row_list

        for i in range(0, len(rowname)):
            c = indexsummary.cell(row=rowlist[i], column=1)
            c.value = rowname[i]
            c.font = datafont

        # 填入整数型数据
        data_name = ['volumn', 'east_volumn', 'mid_volumn', 'west_volumn',
                     'under_90_volumn', '90_144_volumn', 'above_144_volumn',
                     'DONGBEI_volumn', 'HUABEI_volumn', 'HUADONG_volumn', 'HUAZHONG_volumn', 'HUANAN_volumn', 'XINAN_volumn', 'XIBEI_volumn', 
                     'firstline_volumn', 'secondline_volumn', 'thirdline_volumn', 'forth_volumn', 'ZSJ_volumn', 'CSJ_volumn', 'HBH_volumn']
        for i in range(0, len(int_row_list)):
            for j in range(2, col_max):
                c = indexsummary.cell(row=int_row_list[i], column=j)
                c.value = kwargs[data_name[i]][j - 2]
                c.font = datafont

        # 填入浮点数型数据
        data_name = ['index', 'east_index', 'mid_index', 'west_index',
                     'index_under_90', 'index_90_144', 'index_above_144',
                     'DONGBEI_index', 'HUABEI_index', 'HUADONG_index', 'HUAZHONG_index', 'HUANAN_index', 'XINAN_index', 'XIBEI_index', 
                     'firstline_index', 'secondline_index', 'thirdline_index', 'forth_index', 'ZSJ_index', 'CSJ_index', 'HBH_index']
        for i in range(0, len(float_row_list)):
            for j in range(2, col_max):
                c = indexsummary.cell(row=float_row_list[i], column=j)
                c.value = float('%.2f' % kwargs[data_name[i]][j - 2])
                c.font = datafont

        # 填入同比数据
        data_name = ['volumn_year_on_year',
                     'year_on_year', 'east_year_on_year', 'mid_year_on_year', 'west_year_on_year',
                     'year_on_year_under_90', 'year_on_year_90_144', 'year_on_year_above_144',
                     'DONGBEI_year_on_year', 'HUABEI_year_on_year', 'HUADONG_year_on_year', 'HUAZHONG_year_on_year', 'HUANAN_year_on_year', 'XINAN_year_on_year', 'XIBEI_year_on_year', 
                     'firstline_year_on_year', 'secondline_year_on_year', 'thirdline_year_on_year', 'forth_year_on_year', 'ZSJ_year_on_year', 'CSJ_year_on_year', 'HBH_year_on_year']
        for i in range(0, len(year_row_list)):
            # print(len(kwargs[data_name[i]]))
            for j in range(2, 14):
                c = indexsummary.cell(row=year_row_list[i], column=j)
                c.value = '——'
                c.font = datafont
            for j in range(14, col_max):
                c = indexsummary.cell(row=year_row_list[i], column=j)
                c.value = '{:.2%}'.format(kwargs[data_name[i]][j - 14])
                c.font = datafont

        # 填入环比数据
        data_name = ['volumn_chain',
                     'chain', 'east_chain', 'mid_chain', 'west_chain',
                     'chain_under_90', 'chain_90_144', 'chain_above_144',
                     'DONGBEI_chain', 'HUABEI_chain', 'HUADONG_chain', 'HUAZHONG_chain', 'HUANAN_chain', 'XINAN_chain', 'XIBEI_chain', 
                     'firstline_chain', 'secondline_chain', 'thirdline_chain', 'forth_chain', 'ZSJ_chain', 'CSJ_chain', 'HBH_chain']
        for i in range(0, len(year_row_list)):
            c = indexsummary.cell(row=chain_row_list[i], column=2)
            c.value = '——'
            c.font = datafont
            for j in range(3, col_max):
                c = indexsummary.cell(row=chain_row_list[i], column=j)
                c.value = '{:.2%}'.format(kwargs[data_name[i]][j - 3])
                c.font = datafont
    
    def RadioPlot(self, **kwargs):
        # 参数列表:['index','year_on_year','chain','year_on_year_plot','chain_plot']

        # 表格基本格式设置
        radioplot = self.report['全国指数同比环比图']
        titlefont = Font(u'宋体', size=12, bold=True, color='000000')
        datafont = Font(u'宋体', size=10, bold=False, color='000000')
        bluefill = PatternFill("solid", fgColor='FFC0D9D9')
        rowmax = len(kwargs['index']) + 2

        # 填充颜色
        for i in range(1, 5):
            c = radioplot.cell(row=1, column=i)
            c.fill = bluefill

        # 表格边框加粗
        for i in range(1, rowmax):
            for j in range(1, 5):
                c = radioplot.cell(row=i, column=j)
                c.border = self.border

        # 填入表头
        c = radioplot.cell(row=1, column=2)
        c.value = '指数'
        c.font = titlefont
        c = radioplot.cell(row=1, column=3)
        c.value = '环比'
        c.font = titlefont
        c = radioplot.cell(row=1, column=4)
        c.value = '同比'
        c.font = titlefont

        # 填入日期
        for i in range(2, rowmax):
            c = radioplot.cell(row=i, column=1)
            c.value = self.to_date(i - 2)
            c.font = titlefont

        # 填入数据
        for i in range(2, rowmax):
            c = radioplot.cell(row=i, column=2)
            c.value = float('%.2f' % kwargs['index'][i - 2])
            c.font = datafont

            c = radioplot.cell(row=i, column=3)
            if i < 3:
                c.value = '——'
            else:
                c.value = float('%.2f' % (kwargs['chain'][i - 3] * 100))
            c.font = datafont

            c = radioplot.cell(row=i, column=4)
            if i < 14:
                c.value = '——'
            else:
                c.value = float('%.2f' % (kwargs['year_on_year'][i - 14] * 100))
            c.font = datafont

        # 添加图片
        img = Image(kwargs['year_on_year_plot'])
        img.width, img.height = 1200, 300
        radioplot.add_image(img, 'E8')

        img = Image(kwargs['chain_plot'])
        img.width, img.height = 1200, 300
        radioplot.add_image(img, 'E27')  
 
    def SimplyPlot(self, **kwargs):
        # 参数列表:['url_index_plot','url_block_plot','url_area_plot', 'url_7area_plot', 'url_line_plot', 'url_s_area_plot']
        simplyplot = self.report['作图（简版）']
        simplyplot.column_dimensions['A'].height = 4.0
        titlefont = Font(u'宋体', size=14, bold=True, color='000000')
        grayfill = PatternFill("solid", fgColor='00C0C0C0')

        fill_line = [(28*i + 1) for i in range(0,6)]
        # 填充颜色
        for j in fill_line:
            for i in range(1, 16):
                c = simplyplot.cell(row=j, column=i)
                c.fill = grayfill
        plot_name = ['全国（简版封面&最简版图1）', '各地区（最简版图2）', '各面积（最简版图3）', 
        '7区域（最简版图4）', '各线城市（最简版图5）', '重点区域（最简版图6）']
        url_list = ['url_index_plot','url_block_plot','url_area_plot', 'url_7area_plot', 'url_line_plot', 'url_s_area_plot']
        # 添加标题、图片，合并单元格
        for i in range(0,6):
            c = simplyplot.cell(row=fill_line[i], column=1)
            c.value = plot_name[i]
            c.font = titlefont
            img = Image(kwargs[url_list[i]])
            img.width, img.height = 1000, 500
            simplyplot.add_image(img, 'A' + str(fill_line[i] + 1))
            simplyplot.merge_cells('A' + str(fill_line[i]) + ':F' + str(fill_line[i]))

    def ComplexPlot(self, **kwargs):
        # 参数列表:['url_index_plot','url_block_plot','url_area_plot', 'url_7area_plot', 'url_line_plot', 'url_s_area_plot']
        complexplot = self.report['作图（详细）']
        complexplot.column_dimensions['A'].height = 4.0
        titlefont = Font(u'宋体', size=14, bold=True, color='000000')
        grayfill = PatternFill("solid", fgColor='00C0C0C0')

        fill_line = [(28*i + 1) for i in range(0,6)]
        # 填充颜色
        for j in fill_line:
            for i in range(1, 25):
                c = complexplot.cell(row=j, column=i)
                c.fill = grayfill
        plot_name = ['全国（详细）', '各地区（详细）', '各面积（详细）', 
        '7区域（详细）', '各线城市（详细）', '重点区域（详细）']
        url_list = ['url_index_plot','url_block_plot','url_area_plot', 'url_7area_plot', 'url_line_plot', 'url_s_area_plot']
        # 添加标题、图片，合并单元格
        for i in range(0,6):
            c = complexplot.cell(row=fill_line[i], column=1)
            c.value = plot_name[i]
            c.font = titlefont
            img = Image(kwargs[url_list[i]])
            img.width, img.height = 1200, 500
            complexplot.add_image(img, 'A' + str(fill_line[i] + 1))
            complexplot.merge_cells('A' + str(fill_line[i]) + ':F' + str(fill_line[i]))

    def CitySummary(self, informationlist: list):
        # list每项的参数列表:{'city_name','index_this_month','index_last_month','index_last_year',
        # 'chain_radio','year_on_year','volumn_chain','volumn_year'}

        # 表格基本格式设置
        citysummary = self.report['城市汇总']
        citysummary.freeze_panes = 'C2'
        citysummary.column_dimensions['A'].width = 5.0
        datafont = Font(u'宋体', size=10, bold=False, color='000000')
        col_max = 10
        row_max = len(informationlist) + 2

        # 表格边框加粗
        for i in range(1, row_max):
            for j in range(1, col_max):
                c = citysummary.cell(row=i, column=j)
                c.border = self.border

        # 填入表头
        col_name_list = ['当月指数值', '上月指数值', '去年同月指数值',
                         '环比', '同比', '交易量环比', '交易量同比']
        for i in range(3, col_max):
            c = citysummary.cell(row=1, column=i)
            c.value = col_name_list[i - 3]
            c.font = datafont

        # 填入数据
        col_data_list = ['city_name', 'index_this_month', 'index_last_month', 'index_last_year',
                         'chain_radio', 'year_on_year', 'volumn_chain', 'volumn_year']
        for i in range(2, row_max):
            c = citysummary.cell(row=i, column=1)
            c.value = i - 1
            c.font = datafont
            for j in range(2, col_max):
                c = citysummary.cell(row=i, column=j)
                if j == 2:
                    c.value = informationlist[i - 2]['city_name']
                elif j < 6:
                    c.value = float('%.2f' % informationlist[i - 2][col_data_list[j - 2]])
                else:
                    c.value = '%.2f' % (informationlist[i - 2][col_data_list[j - 2]] * 100) + '%'
                c.font = datafont
    
    def LineSummary(self, linelist: list):
        # list每个项的参数：{'index', 'chain', 'year_on_year'}
        linesummary = self.report['各线城市汇总']
        titlefont = Font(u'宋体', size=12, bold=True, color='000000')
        datafont = Font(u'宋体', size=10, bold=False, color='000000')
        # 填入表头
        c = linesummary.cell(row=1, column=2)
        c.value = '当月指数值'
        c.font = titlefont

        c = linesummary.cell(row=1, column=3)
        c.value = '环比'
        c.font = titlefont

        c = linesummary.cell(row=1, column=4)
        c.value = '同比'
        c.font = titlefont

        c = linesummary.cell(row=2, column=1)
        c.value = '一线城市'
        c.font = datafont

        c = linesummary.cell(row=3, column=1)
        c.value = '二线城市'
        c.font = datafont

        c = linesummary.cell(row=4, column=1)
        c.value = '三线城市'
        c.font = datafont

        c = linesummary.cell(row=5, column=1)
        c.value = '四线城市'
        c.font = datafont

        # 填入数据
        for i in range(0,4):
            c = linesummary.cell(row=2 + i, column=2)
            c.value = float('%.3f' % linelist[i]['index'])
            c.font = datafont
            c = linesummary.cell(row=2 + i, column=3)
            c.value = '{:.2%}'.format(linelist[i]['chain'])
            c.font = datafont
            c = linesummary.cell(row=2 + i, column=4)
            c.value = '{:.2%}'.format(linelist[i]['year_on_year'])
            c.font = datafont

    def CitySort(self, month: int, cityinformation: list):
        # 列表中每个元素应有参数:{'city_name','chain0','chain1' ,'chain2',
        # 'chain3' 'year0' ,'year1' ,'year2','year3'}

        # 表格基本格式设置
        citysort = self.report['城市排序']
        datafont = Font(u'宋体', size=10, bold=True, color='000000')
        Reddatafont = Font(u'宋体', size=10, bold=False, color='FF0000')
        bluefill = PatternFill("solid", fgColor='FF3299CC')
        rowmax = len(cityinformation) + 3
        colmax = 11
        align = Alignment(horizontal='center', vertical='center')

        # 表格边框加粗
        for i in range(1, rowmax):
            for j in range(1, colmax):
                c = citysort.cell(row=i, column=j)
                c.border = self.border

        # 颜色填充
        for j in range(1, colmax):
            for i in [1, 2]:
                c = citysort.cell(row=i, column=j)
                c.fill = bluefill

        # 添加表头
        c = citysort.cell(row=1, column=3)
        c.value = '环比'
        c.font = datafont
        c.alignment = align
        citysort.merge_cells('C1:F1')

        c = citysort.cell(row=1, column=7)
        c.value = '同比'
        c.font = datafont
        c.alignment = align
        citysort.merge_cells('G1:J1')

        c = citysort.cell(row=2, column=1)
        c.value = '编号'
        c.font = datafont

        c = citysort.cell(row=2, column=2)
        c.value = '城市'
        c.font = datafont

        for i in range(0, 4):
            c = citysort.cell(row=2, column=3 + i)
            if month - i > 0:
                c.value = str(month - i) + '月'
            else:
                c.value = str(12 + month - i) + '月'
            c.font = datafont
            c = citysort.cell(row=2, column=7 + i)
            if month - i > 0:
                c.value = str(month - i) + '月'
            else:
                c.value = str(12 + month - i) + '月'
            c.font = datafont

        # 填入数据
        col_data_list = ['city_name', 'chain0', 'chain1', 'chain2', 'chain3', 'year0', 'year1', 'year2', 'year3']
        for i in range(3, rowmax):
            c = citysort.cell(row=i, column=1)
            c.value = i - 2
            c.font = datafont
            for j in range(2, colmax):
                c = citysort.cell(row=i, column=j)
                if j == 2:
                    c.value = cityinformation[i - 3]['city_name']
                else:
                    c.value = float('%.2f' % (cityinformation[i - 3][col_data_list[j - 2]] * 100))
                    if c.value == 0 or c.value == -100:
                        c.value = ''
                    elif c.value < 0:
                        c.font = Reddatafont

        # 添加排序功能
        citysort.auto_filter.ref = "A2:J" + str(colmax - 1)
        citysort.auto_filter.add_sort_condition("A3:A" + str(colmax - 1))

    def CityIndex(self, citylist: list, **kwargs):
        # 参数列表citylist中的城市的数据

        # 表格基本格式设置
        cityindex = self.report['各城市指数']
        cityindex.freeze_panes = 'D3'
        cityindex.column_dimensions['A'].width = 5.0
        datafont = Font(u'宋体', size=10, bold=False, color='000000')
        col_max = len(kwargs[citylist[0]]) + 3
        row_max = len(citylist) + 2

        # 填入表头时间
        for i in range(3, col_max):
            c = cityindex.cell(row=1, column=i)
            c.value = self.to_date(i - 3)
            c.font = datafont
        # 填入序号
        for i in range(2, row_max):
            c = cityindex.cell(row=i, column=1)
            c.value = i - 1
            c.font = datafont
            c = cityindex.cell(row=i, column=2)
            c.value = citylist[i - 2]
            c.font = datafont
        # 填入数据
        for i in range(2, row_max):
            for j in range(3, col_max):
                c = cityindex.cell(row=i, column=j)
                c.value = float('%.2f' % kwargs[citylist[i - 2]][j - 3])
                c.font = datafont

    def CityVolumn(self, citylist: list, **kwargs):
        # 参数列表citylist中的城市的数据

        # 表格基本格式设置
        cityvolumn = self.report['各城市样本量']
        cityvolumn.freeze_panes = 'D3'
        cityvolumn.column_dimensions['A'].width = 5.0
        datafont = Font(u'宋体', size=10, bold=False, color='000000')
        col_max = len(kwargs[citylist[0]]) + 3
        row_max = len(citylist) + 2

        # 填入表头时间
        for i in range(3, col_max):
            c = cityvolumn.cell(row=1, column=i)
            c.value = self.to_date(i - 3)
            c.font = datafont
        # 填入序号
        for i in range(2, row_max):
            c = cityvolumn.cell(row=i, column=1)
            c.value = i - 1
            c.font = datafont
            c = cityvolumn.cell(row=i, column=2)
            c.value = citylist[i - 2]
            c.font = datafont
        # 填入数据
        for i in range(2, row_max):
            for j in range(3, col_max):
                c = cityvolumn.cell(row=i, column=j)
                c.value = kwargs[citylist[i - 2]][j - 3]
                c.font = datafont

    def CityPlot(self, citylist: list, poltlist: list):
        cityindex = self.report['各城市作图']
        datafont = Font(u'宋体', size=10, bold=False, color='000000')

        # 25
        for i in range(0, len(citylist)):
            c = cityindex.cell(row=25 * i + 1, column=1)
            c.value = '图' + str(i + 1) + ' ' + citylist[i] + '\"城房指数\"'
            c.font = datafont
            img = Image(poltlist[i])
            img.width, img.height = 1600, 400
            cityindex.add_image(img, 'A' + str(25 * i + 2))
    
    def LineCity(self, citylist: list, linelist: list):
        # citylist中参数: {'name', 'line', 'index', 'chain', 'year_on_year'}
        # linelist中参数: {'index', 'chain', 'year_on_year'}
        linecity = self.report['各线城市子市场']
        titlefont = Font(u'宋体', size=12, bold=True, color='000000')
        datafont = Font(u'宋体', size=10, bold=False, color='000000')
        city_line = [[],[],[],[]]
        line_name = ['一线城市','二线城市','三线城市','四线城市']
        for city in citylist:
                city_line[city['line'] - 1].append(city)
        name_list = ['各线城市子市场','指数值','同比','环比']
        key_list = ['name' ,'index', 'chain', 'year_on_year']
        for i in range(0,4):
            c = linecity.cell(row=1, column=i + 1)
            c.value = name_list[i]
            c.font = titlefont
        row_num = 2
        for i in range(0,4):
            c = linecity.cell(row=row_num, column=1)
            c.value = line_name[i]
            c.font = titlefont
            c = linecity.cell(row=row_num, column=2)
            c.value = '%.4f' % linelist[i]['index']
            c.font = titlefont
            c = linecity.cell(row=row_num, column=3)
            c.value = '{:.2%}'.format(linelist[i]['chain'])
            c.font = titlefont
            c = linecity.cell(row=row_num, column=4)
            c.value = '{:.2%}'.format(linelist[i]['year_on_year'])
            c.font = titlefont
            row_num += 1
            for city in city_line[i]:
                c = linecity.cell(row=row_num, column=1)
                c.value = city['name']
                c.font = datafont
                c = linecity.cell(row=row_num, column=2)
                c.value = '%.4f' % city['index']
                c.font = datafont
                c = linecity.cell(row=row_num, column=3)
                c.value = '{:.2%}'.format(city['chain'])
                c.font = datafont
                c = linecity.cell(row=row_num, column=4)
                c.value = '{:.2%}'.format(city['year_on_year'])
                c.font = datafont
                row_num += 1



class GenWordReport90:
    def __init__(self, url):
        self.url = url
        self.doc = Document()

    def EndReport(self):
        self.doc.save(self.url)

    def Firstpage(self, year: int, month: int, **kwargs):
        title = str(year) + '年' + str(month) + '月全国90个重点城市“城房指数”月报'
        self.doc.add_paragraph('')
        self.doc.add_paragraph('')
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run(title)
        tit.font.name = '黑体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        tit.font.size = Pt(28)
        tit.font.bold = True

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('（试行）')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(20)
        self.doc.add_paragraph('')

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('全国90个重点城市各月“城房指数”汇总值')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        p = self.doc.add_paragraph()
        pic = p.add_run().add_picture(kwargs['index_image_url_90'])
        pic.height = Cm(10.96)
        pic.width = Cm(17.65)

        self.doc.add_section(start_type=None)

    def Secondpage(self, year: int, month: int, citylist: list, **kwargs):
        # 参数列表:['index_90', 'chain_90', 'year_on_year_90', 'city_name_90', 'city_chain_90', 'city_year_on_year_90']
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('简要说明')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(20)
        tit.font.bold = True

        p = self.doc.add_paragraph()
        l = len(kwargs['index_90'])
        index = float('%.2f' % kwargs['index_90'][l - 1])
        chain = float('%.2f' % (float(kwargs['chain_90'][l - 1]) * 100))
        year_on_year = float('%.2f' % (float(kwargs['year_on_year_90'][l - 1]) * 100))
        s = 0
        if chain < 0:
            str1 = '下降' + str(-chain) + '%。'
        else:
            str1 = '上涨' + str(chain) + '%。'
        if year_on_year < 0:
            str2 = '下降' + str(-year_on_year) + '%'
        else:
            str2 = '上涨' + str(year_on_year) + '%'
        zero = []
        for i in range(len(citylist)):
            if float(kwargs['city_index_90'][i]) == 0:
                zero.append(i)
        if len(zero) == 0:
            add = p.add_run(
                '  ' + str(year) + '年' + str(month) + '月”城房指数“全国汇总值为' + str(index) + '点，同比' + str2 + '，环比' + str1)
            add.font.name = '仿宋_GB2312'
            add.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            add.font.size = Pt(14)
        else:
            add = p.add_run('    ' + str(year) + '年' + str(month) + '月”城房指数“全国')
            add.font.name = '仿宋_GB2312'
            add.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            add.font.size = Pt(14)

            add = p.add_run('1')
            add.font.superscript = True
            add.font.name = '仿宋_GB2312'
            add.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            add.font.size = Pt(14)

            add = p.add_run('汇总值为' + str(index) + '点，同比' + str2 + '，环比' + str1)
            add.font.name = '仿宋_GB2312'
            add.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            add.font.size = Pt(14)
        # 统计环比涨跌

        l = len(citylist)
        higher = []
        num_higher = 0
        higher_than_5 = []
        num_lower = 0
        lower = []
        lower_than_5 = []
        for i in range(l):
            if float(kwargs['city_chain_90'][i]) > 0:
                higher.append(i)
                num_higher += 1
                if float(kwargs['city_chain_90'][i]) > 0.05:
                    higher_than_5.append(i)
            elif (float(kwargs['city_chain_90'][i]) < 0) & (float(kwargs['city_chain_90'][i]) != -1):
                lower.append(i)
                num_lower += 1
                if float(kwargs['city_chain_90'][i]) < -0.05:
                    lower_than_5.append(i)
        # 环比涨幅

        p = self.doc.add_paragraph()
        if num_higher == 0:
            tit = p.add_run('    没有任何一个城市的“城房指数”环比上涨；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('    有' + str(num_higher) + '个城市的“城房指数”环比上涨，')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

            if len(higher_than_5) == 0:
                tit = p.add_run('但涨幅均在5%以内；')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
            else:
                tit = p.add_run('其中')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)

                for i in range(len(higher_than_5)):
                    s += 1
                    str1 = '(' + str(float('%.2f' % (float(kwargs['city_chain_90'][higher_than_5[i]]) * 100)))
                    if i != (len(higher_than_5) - 1):
                        tit = p.add_run(citylist[higher_than_5[i]])
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str(s))
                        tit.font.superscript = True
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str1 + '%）、')
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)
                    else:
                        tit = p.add_run(citylist[higher_than_5[i]])
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str(s))
                        tit.font.superscript = True
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str1 + '%）')
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                tit = p.add_run(str(len(higher_than_5)) + '个城市涨幅超过5%；')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
        # 环比跌幅

        if num_lower == 0:
            tit = p.add_run('    没有任何一个城市的“城房指数”环比下降；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('有' + str(num_lower) + '个城市的“城房指数”环比下降，')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

            if len(lower_than_5) == 0:
                tit = p.add_run('但降幅均在5%以内；')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
            else:
                tit = p.add_run('其中')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)

                for i in range(len(lower_than_5)):
                    s += 1
                    str1 = '(' + str(float('%.2f' % (float(kwargs['city_chain_90'][lower_than_5[i]]) * 100)))
                    if i != (len(lower_than_5) - 1):
                        tit = p.add_run(citylist[lower_than_5[i]])
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str(s))
                        tit.font.superscript = True
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str1 + '%）、')
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)
                    else:
                        tit = p.add_run(citylist[lower_than_5[i]])
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str(s))
                        tit.font.superscript = True
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                        tit = p.add_run(str1 + '%）')
                        tit.font.name = '仿宋_GB2312'
                        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        tit.font.size = Pt(14)

                tit = p.add_run(str(len(lower_than_5)) + '个城市降幅超过5%。')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
        # 统计同比涨降

        higher = []
        num_higher = 0
        higher_than_15 = []
        num_lower = 0
        lower = []
        lower_than_15 = []
        for i in range(l):
            if float(kwargs['city_year_on_year_90'][i]) > 0:
                higher.append(i)
                num_higher += 1
                if float(kwargs['city_year_on_year_90'][i]) > 0.15:
                    higher_than_15.append(i)
            elif (float(kwargs['city_year_on_year_90'][i]) < 0) & (float(kwargs['city_year_on_year_90'][i]) != -1):
                lower.append(i)
                num_lower += 1
                if float(kwargs['city_year_on_year_90'][i]) < -0.15:
                    lower_than_15.append(i)
        # 同比涨幅

        p = self.doc.add_paragraph()
        if num_higher == 0:
            tit = p.add_run('    没有任何一个城市的“城房指数”同比上涨；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('    有' + str(num_higher) + '个城市的“城房指数”同比上涨，')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

            if len(higher_than_15) == 0:
                tit = p.add_run('但涨幅均在15%以内；')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
            else:
                tit = p.add_run('其中')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)

                for i in range(len(higher_than_15)):
                    str1 = '(' + str(float('%.2f' % (float(kwargs['city_year_on_year_90'][higher_than_15[i]]) * 100)))
                    if i != (len(higher_than_15) - 1):
                        tit = p.add_run(citylist[higher_than_15[i]] + str1 + '%）、')
                    else:
                        tit = p.add_run(citylist[higher_than_15[i]] + str1 + '%）')
                    tit.font.name = '仿宋_GB2312'
                    tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    tit.font.size = Pt(14)

                tit = p.add_run(str(len(higher_than_15)) + '个城市涨幅超过15%；')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
        # 同比跌幅

        if num_lower == 0:
            tit = p.add_run('    没有任何一个城市的“城房指数”同比下降；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('有' + str(num_lower) + '个城市的“城房指数”同比下降，')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

            if len(lower_than_15) == 0:
                tit = p.add_run('但降幅均在15%以内；')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
            else:
                tit = p.add_run('其中')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)

                for i in range(len(lower_than_15)):
                    str1 = '(' + str(float('%.2f' % (float(kwargs['city_year_on_year_90'][lower_than_15[i]]) * 100)))
                    if i != (len(lower_than_15) - 1):
                        tit = p.add_run(citylist[lower_than_15[i]] + str1 + '%）、')
                    else:
                        tit = p.add_run(citylist[lower_than_15[i]] + str1 + '%）')
                    tit.font.name = '仿宋_GB2312'
                    tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    tit.font.size = Pt(14)

                tit = p.add_run(str(len(lower_than_15)) + '个城市降幅超过15%。')
                tit.font.name = '仿宋_GB2312'
                tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                tit.font.size = Pt(14)
        # 插入图片

        p = self.doc.add_paragraph()
        pic = p.add_run().add_picture(kwargs['yearonyear_image_url_90'])
        pic.height = Cm(5.46)
        pic.width = Cm(14.67)

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('图1 90城市“城房指数”同比变化图')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        # 脚注
        p = self.doc.add_paragraph()
        tit = p.add_run('注：')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(9)
        tit.font.bold = True

        # 脚注正文
        for i in range(len(lower_than_5)):
            p = self.doc.add_paragraph()
            tit = p.add_run(str(i + 1) + citylist[lower_than_5[i]] + '(' + str(float('%.2f' % (float(
                kwargs['city_chain_90'][
                    lower_than_5[i]]) * 100))) + '):本月均价变化率为 ，其中由于结构性因素导致的变化有 。本月均价 幅度 ，结构性因素 幅度 ，反映出同质性价格 。')
            tit.font.name = '宋体'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            tit.font.size = Pt(9)

        self.doc.add_page_break()

        p = self.doc.add_paragraph()
        pic = p.add_run().add_picture(kwargs['chain_image_url_90'])
        pic.height = Cm(5.23)
        pic.width = Cm(14.65)

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('图2 90城市“城房指数”环比变化图')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        self.doc.add_section(start_type=None)

    def Charts(self, year: int, month: int, citylist: list, **kwargs):

        # 数据从2009年1月开始
        # 第一页45行；其余页48行
        p = self.doc.add_paragraph()
        section = self.doc.sections[2]
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        if int(((year - 2009) * 12 + month) // 2) == 0:
            row_num = int(((year - 2009) * 12 + month) // 2)
        else:
            row_num = int(((year - 2009) * 12 + month) // 2) + 1
        r_num = row_num
        row_num = row_num + int(row_num - 45) // 47 + 2
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('表1 全国90个重点城市“城房指数”各月汇总值')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        # 表1
        table = self.doc.add_table(row_num, 8, 'Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        data_row = 1
        for i in range(row_num):
            if ((i == 0) or ((i - 45) % 48 == 0)):
                p = table.cell(i, 0).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'月份')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 1).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'指数值')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 2).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'环比')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 3).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'同比')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 4).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'月份')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 5).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'指数值')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 6).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'环比')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 7).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'同比')
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True


            else:
                p = table.cell(i, 0).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                m = str(int(data_row) % 12)
                if m == '0':
                    m = '12'
                y = (int(data_row) // 12) + 2008
                if (int(data_row) % 12) == 0:
                    y -= 1
                y = str(y)
                str1 = y + '年' + m + '月'
                run = p.add_run(u'%s' % str1)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 1).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % kwargs['index_90'][data_row - 1]))
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 2).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % (float(kwargs['chain_90'][data_row - 1]) * 100))) + "%"
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                p = table.cell(i, 3).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(float('%.2f' % (float(kwargs['year_on_year_90'][data_row - 1]) * 100))) + "%"
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)

                if i != row_num - 1:
                    p = table.cell(i, 4).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    m = str(int(data_row + r_num) % 12)
                    if m == '0':
                        m = '12'
                    y = (int(data_row + r_num) // 12) + 2008
                    if (int(data_row + r_num) % 12) == 0:
                        y -= 1
                    y = str(y)
                    str1 = y + '年' + m + '月'
                    run = p.add_run(u'%s' % str1)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                    p = table.cell(i, 5).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(float('%.2f' % kwargs['index_90'][data_row + r_num - 1]))
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                    p = table.cell(i, 6).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(float('%.2f' % (float(kwargs['chain_90'][data_row + r_num - 1]) * 100))) + "%"
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                    p = table.cell(i, 7).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(float('%.2f' % (float(kwargs['year_on_year_90'][data_row + r_num - 1]) * 100))) + "%"
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                data_row += 1

        self.doc.add_page_break()

        # 表2
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('表2 %s年%s月90个重点城市“城房指数”' % (str(year), str(month)))
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        table = self.doc.add_table(47, 8, 'Table Grid')

        # kwargs['city_index_90']:当月城市“城房指数”值的列表，按城市顺序排列，'city_chain_90'以及'city_year_on_year_90'以此类推
        for i in range(47):
            if ((i != 0) and (i != 45)):
                p = table.cell(i, 0).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(citylist[i - 1])
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                if i != 46:
                    if float(kwargs['city_index_90'][i - 1]) != 0:
                        p = table.cell(i, 1).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % kwargs['city_index_90'][i - 1]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 2).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(kwargs['city_chain_90'][i - 1]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 3).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(kwargs['city_year_on_year_90'][i - 1]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                    else:
                        p = table.cell(i, 1).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 2).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 3).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    p = table.cell(i, 4).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(citylist[i + 44])
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                    if float(kwargs['city_index_90'][i + 44]) != 0:
                        p = table.cell(i, 5).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % kwargs['city_index_90'][i + 19]))
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 6).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(kwargs['city_chain_90'][i + 44]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 7).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(kwargs['city_year_on_year_90'][i + 44]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                    else:
                        p = table.cell(i, 5).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 6).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 7).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                else:
                    if float(kwargs['city_index_90'][i - 2]) != 0:
                        p = table.cell(i, 1).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % kwargs['city_index_90'][i - 2]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 2).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(kwargs['city_chain_90'][i - 2]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 3).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(kwargs['city_year_on_year_90'][i - 2]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                    else:
                        p = table.cell(i, 1).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 2).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 3).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    p = table.cell(i, 4).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(citylist[i + 43])
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                    if float(kwargs['city_index_90'][i + 43]) != 0:
                        p = table.cell(i, 5).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % kwargs['city_index_90'][i + 19]))
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 6).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(kwargs['city_chain_90'][i + 43]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 7).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(kwargs['city_year_on_year_90'][i + 43]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                    else:
                        p = table.cell(i, 5).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 6).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 7).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

            else:
                p = table.cell(i, 0).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '城市'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 1).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '指数值'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 2).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '环比'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 3).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '同比'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 4).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '城市'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 5).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '指数值'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 6).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '环比'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 7).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '同比'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

        self.doc.add_page_break()

        # 表3
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('表3 %s年%s月90个重点城市“城房指数”样本量' % (str(year), str(month)))
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        table = self.doc.add_table(47, 8, 'Table Grid')

        # kwargs['city_volume_90']:当月城市“城房指数”值的列表，按城市顺序排列，'city_chain_90'以及'city_year_on_year_90'以此类推
        # kwargs['city_volume_1_90']:上月城市“城房指数”值的列表，按城市顺序排列
        # kwargs['city_volume_2_90']:两个月前城市“城房指数”值的列表，按城市顺序排列
        for i in range(47):
            if (i != 0) and (i != 45):
                if i != 46:
                    p = table.cell(i, 0).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(citylist[i - 1])
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                    if int(kwargs['city_volume_2_90'][i - 1]) != 0:
                        p = table.cell(i, 1).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(int(kwargs['city_volume_2_90'][i - 1]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    if int(kwargs['city_volume_1_90'][i - 1]) != 0:
                        p = table.cell(i, 2).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(int(kwargs['city_volume_1_90'][i - 1]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    if int(kwargs['city_volume_90'][i - 1]) != 0:
                        p = table.cell(i, 3).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(int(kwargs['city_volume_90'][i - 1]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    p = table.cell(i, 4).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(citylist[i + 19])
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                    if int(kwargs['city_volume_2_90'][i + 19]) != 0:
                        p = table.cell(i, 5).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(int(kwargs['city_volume_2_90'][i + 19]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    if int(kwargs['city_volume_1_90'][i + 19]) != 0:
                        p = table.cell(i, 6).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(int(kwargs['city_volume_1_90'][i + 19]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    if int(kwargs['city_volume_90'][i + 19]) != 0:
                        p = table.cell(i, 7).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(int(kwargs['city_volume_90'][i + 19]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                else:
                    p = table.cell(i, 0).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(citylist[i - 2])
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                    if int(kwargs['city_volume_2_90'][i - 2]) != 0:
                        p = table.cell(i, 1).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(int(kwargs['city_volume_2_90'][i - 1]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    if int(kwargs['city_volume_1_90'][i - 2]) != 0:
                        p = table.cell(i, 2).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(int(kwargs['city_volume_1_90'][i - 1]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    if int(kwargs['city_volume_90'][i - 2]) != 0:
                        p = table.cell(i, 3).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(int(kwargs['city_volume_90'][i - 1]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    p = table.cell(i, 4).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(citylist[i + 43])
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                    if int(kwargs['city_volume_2_90'][i + 43]) != 0:
                        p = table.cell(i, 5).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(int(kwargs['city_volume_2_90'][i + 19]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    if int(kwargs['city_volume_1_90'][i + 43]) != 0:
                        p = table.cell(i, 6).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(int(kwargs['city_volume_1_90'][i + 19]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    if int(kwargs['city_volume_90'][i + 43]) != 0:
                        p = table.cell(i, 7).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(int(kwargs['city_volume_90'][i + 19]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

            else:
                if int(month) == 2:
                    mon = [12, 1, 2]
                elif int(month) == 1:
                    mon = [11, 12, 1]
                else:
                    mon = [int(month) - 2, int(month) - 1, int(month)]

                p = table.cell(i, 0).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '城市'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 1).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(mon[0]) + '月样本量'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 2).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(mon[1]) + '月样本量'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 3).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(mon[2]) + '月样本量'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 4).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '城市'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 5).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(mon[0]) + '月样本量'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 6).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(mon[1]) + '月样本量'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 7).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = str(mon[2]) + '月样本量'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True
        self.doc.add_section(start_type=None)

        # 表4
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('表4 %s年%s月90个重点城市“城房指数”各线城市子市场' % (str(year), str(month)))
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        table = self.doc.add_table(49, 8, 'Table Grid')

        # kwargs['city_index_90_1']:一线城市“城房指数”的列表，按城市顺序排列，第一项为汇总的指数值
        # 'city_chain_90_1'以及'city_year_on_year_90_1'以此类推
        # kwargs['cityname_1']:一线城市名，第一项为“一线城市”，以此类推
        citylist_4 = []
        city_index = []
        city_chain = []
        city_year_on_year = []

        for i in range(len(kwargs['city_index_90_1'])):
            citylist_4.append(kwargs['cityname_1'][i])
            city_index.append(kwargs['city_index_90_1'][i])
            city_chain.append(kwargs['city_chain_90_1'][i])
            city_year_on_year.append(kwargs['city_year_on_year_90_1'][i])

        for i in range(len(kwargs['city_index_90_2'])):
            citylist_4.append(kwargs['cityname_2'][i])
            city_index.append(kwargs['city_index_90_2'][i])
            city_chain.append(kwargs['city_chain_90_2'][i])
            city_year_on_year.append(kwargs['city_year_on_year_90_2'][i])

        for i in range(len(kwargs['city_index_90_3'])):
            citylist_4.append(kwargs['cityname_3'][i])
            city_index.append(kwargs['city_index_90_3'][i])
            city_chain.append(kwargs['city_chain_90_3'][i])
            city_year_on_year.append(kwargs['city_year_on_year_90_3'][i])

        for i in range(len(kwargs['city_index_90_4'])):
            citylist_4.append(kwargs['cityname_4'][i])
            city_index.append(kwargs['city_index_90_4'][i])
            city_chain.append(kwargs['city_chain_90_4'][i])
            city_year_on_year.append(kwargs['city_year_on_year_90_4'][i])

        for i in range(49):
            if ((i != 0) and (i != 45)):
                if i < 46:
                    b = False
                    p = table.cell(i, 0).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(citylist_4[i - 1])
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)
                    if '城市' in citylist_4[i - 1]:
                        run.font.bold = True
                        b = True
                    if float(city_index[i - 1]) != 0:
                        p = table.cell(i, 1).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % city_index[i - 1]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                        if b == True:
                            run.font.bold = True

                        p = table.cell(i, 2).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(city_chain[i - 1]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                        if b == True:
                            run.font.bold = True

                        p = table.cell(i, 3).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(city_year_on_year[i - 1]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                        if b == True:
                            run.font.bold = True

                    else:
                        p = table.cell(i, 1).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 2).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 3).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    p = table.cell(i, 4).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(citylist[i + 44])
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)

                    if float(kwargs['city_index_90'][i + 44]) != 0:
                        p = table.cell(i, 5).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % kwargs['city_index_90'][i + 19]))
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                        if b == True:
                            run.font.bold = True

                        p = table.cell(i, 6).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(kwargs['city_chain_90'][i + 44]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                        if b == True:
                            run.font.bold = True

                        p = table.cell(i, 7).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(kwargs['city_year_on_year_90'][i + 44]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                        if b == True:
                            run.font.bold = True

                    else:
                        p = table.cell(i, 5).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 6).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 7).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                else:

                    p = table.cell(i, 0).paragraphs[0]
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    data = str(citylist_4[i - 2])
                    run = p.add_run(u'%s' % data)
                    run.font.name = '仿宋_GB2312'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(10)
                    if '城市' in citylist_4[i - 2]:
                        run.font.bold = True
                        b = True

                    if float(city_index[i - 2]) != 0:
                        p = table.cell(i, 1).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % city_index[i - 2]))
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                        if b == True:
                            run.font.bold = True

                        p = table.cell(i, 2).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(city_chain[i - 2]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                        if b == True:
                            run.font.bold = True

                        p = table.cell(i, 3).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(city_year_on_year[i - 2]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                        if b == True:
                            run.font.bold = True
                    else:
                        p = table.cell(i, 1).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 2).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 3).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 4).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(citylist[i + 43])
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                    if float(city_index[i + 43]) != 0:
                        p = table.cell(i, 5).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % city_index[i + 19]))
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                        if b == True:
                            run.font.bold = True

                        p = table.cell(i, 6).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(city_chain[i + 43]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                        if b == True:
                            run.font.bold = True

                        p = table.cell(i, 7).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        data = str(float('%.2f' % (float(city_year_on_year[i + 43]) * 100))) + "%"
                        run = p.add_run(u'%s' % data)
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)
                        if b == True:
                            run.font.bold = True

                    else:
                        p = table.cell(i, 5).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 6).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

                        p = table.cell(i, 7).paragraphs[0]
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(u'————')
                        run.font.name = '仿宋_GB2312'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10)

            else:
                p = table.cell(i, 0).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '各线城市子市场'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 1).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '指数值'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 2).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '环比'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 3).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '同比'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 4).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '各线城市子市场'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 5).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '指数值'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 6).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '环比'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

                p = table.cell(i, 7).paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                data = '同比'
                run = p.add_run(u'%s' % data)
                run.font.name = '仿宋_GB2312'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(10)
                run.font.bold = True

        self.doc.add_page_break()

    def Attach(self, year: int, month: int, citylist: list, **kwargs):

        section = self.doc.sections[3]
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('附：' + str(year) + '年' + str(month) + '月各子市场“城房指数”变化情况')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        # 一
        p = self.doc.add_paragraph()
        tit = p.add_run('一、东中西区域子市场')
        tit.font.name = '黑体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        tit.font.size = Pt(14)

        p = self.doc.add_paragraph()
        tit = p.add_run('    东中西区域子市场中，东部地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['east_index_90'][len(kwargs['east_index_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['east_year_on_year_90'][len(kwargs['east_index_90']) - 1]) > 0:
            tit = p.add_run('上涨' + str(
                float(
                    '%.2f' % (float(kwargs['east_year_on_year_90'][len(kwargs['east_index_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float(
                    '%.2f' % (float(
                        -kwargs['east_year_on_year_90'][len(kwargs['east_index_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['east_chain_90'][len(kwargs['east_index_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(
                    float('%.2f' % (float(kwargs['east_chain_90'][len(kwargs['east_index_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float('%.2f' % (-float(kwargs['east_chain_90'][len(kwargs['east_index_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        tit = p.add_run('中部地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['mid_index_90'][len(kwargs['mid_index_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['mid_year_on_year_90'][len(kwargs['mid_index_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(
                    float('%.2f' % (
                                float(kwargs['mid_year_on_year_90'][len(kwargs['mid_index_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float('%.2f' % (-float(
                        kwargs['mid_year_on_year_90'][len(kwargs['mid_index_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['mid_chain_90'][len(kwargs['mid_index_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(
                    float('%.2f' % (float(kwargs['mid_chain_90'][len(kwargs['mid_index_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float('%.2f' % (-float(kwargs['mid_chain_90'][len(kwargs['mid_index_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        tit = p.add_run('西部地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['west_index_90'][len(kwargs['west_index_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['west_year_on_year_90'][len(kwargs['west_index_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(
                    float('%.2f' % (float(
                        kwargs['west_year_on_year_90'][len(kwargs['west_index_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float(
                    '%.2f' % (-float(
                        kwargs['west_year_on_year_90'][len(kwargs['west_index_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['west_chain_90'][len(kwargs['west_index_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(
                    float('%.2f' % (float(kwargs['west_chain_90'][len(kwargs['west_index_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float('%.2f' % (-float(kwargs['west_chain_90'][len(kwargs['west_index_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        p = self.doc.add_paragraph()
        pic = p.add_run().add_picture(kwargs['index_by_block_image_url_90_90'])
        pic.height = Cm(6.17)
        pic.width = Cm(14.63)

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('附图1 东中西区域子市场走势图')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        # 二
        p = self.doc.add_paragraph()
        tit = p.add_run('二、各面积子市场')
        tit.font.name = '黑体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        tit.font.size = Pt(14)

        p = self.doc.add_paragraph()
        tit = p.add_run('    各面积子市场中，“90平方米以下”指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['index_under_90_90'][len(kwargs['index_under_90_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['year_on_year_under_90_90'][len(kwargs['index_under_90_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (
                        float(
                            kwargs['year_on_year_under_90_90'][len(kwargs['index_under_90_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-float(
                    kwargs['year_on_year_under_90_90'][len(kwargs['index_under_90_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chain_under_90_90'][len(kwargs['index_under_90_90']) - 1]) > 0:
            tit = p.add_run('上涨' + str(
                float(
                    '%.2f' % (float(kwargs['chain_under_90_90'][len(kwargs['index_under_90_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float('%.2f' % (-float(
                        kwargs['chain_under_90_90'][len(kwargs['index_under_90_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        tit = p.add_run('“90至144平方米”指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['index_90_144_90'][len(kwargs['index_90_144_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['year_on_year_90_144_90'][len(kwargs['index_90_144_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float(
                    '%.2f' % (float(
                        kwargs['year_on_year_90_144_90'][len(kwargs['index_90_144_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float(
                    '%.2f' % (-float(
                        kwargs['year_on_year_90_144_90'][len(kwargs['index_90_144_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chain_90_144_90'][len(kwargs['index_90_144_90']) - 1]) > 0:
            tit = p.add_run('上涨' + str(
                float('%.2f' % (float(kwargs['chain_90_144_90'][len(kwargs['index_90_144_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float(
                        '%.2f' % (-float(kwargs['chain_90_144_90'][len(kwargs['index_90_144_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        tit = p.add_run('“144平方米以上”指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(
            str(float('%.2f' % kwargs['index_above_144_90'][len(kwargs['index_above_144_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['year_on_year_above_144_90'][len(kwargs['index_above_144_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (float(
                    kwargs['year_on_year_above_144_90'][len(kwargs['index_above_144_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-float(
                    kwargs['year_on_year_above_144_90'][len(kwargs['index_above_144_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chain_above_144_90'][len(kwargs['index_above_144_90']) - 1]) > 0:
            tit = p.add_run('上涨' + str(
                float('%.2f' % (
                            float(kwargs['chain_above_144_90'][len(kwargs['index_above_144_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('下降' + str(
                float('%.2f' % (
                            -float(kwargs['chain_above_144_90'][len(kwargs['index_above_144_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        p = self.doc.add_paragraph()
        pic = p.add_run().add_picture(kwargs['index_by_buildarea_image_url_90'])
        pic.height = Cm(6.17)
        pic.width = Cm(14.63)

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('附图2 各面积子市场走势图')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        # 三

        p = self.doc.add_paragraph()
        tit = p.add_run('三、七区域子市场')
        tit.font.name = '黑体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        tit.font.size = Pt(14)

        # 东北地区指数
        p = self.doc.add_paragraph()
        tit = p.add_run('    七区域子市场中，东北地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(
            str(float('%.2f' % kwargs['index_northeast_90'][len(kwargs['index_northeast_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['year_on_year_northeast_90'][len(kwargs['index_northeast_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (
                        float(kwargs['year_on_year_northeast_90'][
                                  len(kwargs['index_northeast_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-float(
                    kwargs['year_on_year_northeast_90'][len(kwargs['index_northeast_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chain_northeast_90'][len(kwargs['index_northeast_90']) - 1]) > 0:
            tit = p.add_run('上涨' + str(
                float('%.2f' % (
                            float(kwargs['chain_northeast_90'][len(kwargs['index_northeast_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float('%.2f' % (-float(
                        kwargs['chain_northeast_90'][len(kwargs['index_northeast_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        # 华北地区指数
        tit = p.add_run('华北地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['index_north_90'][len(kwargs['index_north_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['year_on_year_north_90'][len(kwargs['index_north_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float(
                    '%.2f' % (float(
                        kwargs['year_on_year_north_90'][len(kwargs['index_north_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float(
                    '%.2f' % (-float(
                        kwargs['year_on_year_north_90'][len(kwargs['index_north_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chain_north_90'][len(kwargs['index_north_90']) - 1]) > 0:
            tit = p.add_run('上涨' + str(
                float('%.2f' % (float(kwargs['chain_north_90'][len(kwargs['index_north_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float('%.2f' % (-float(kwargs['chain_north_90'][len(kwargs['index_north_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        # 华东地区指数
        tit = p.add_run('华东地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['index_east_90'][len(kwargs['index_east_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['year_on_year_east_90'][len(kwargs['index_east_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (float(
                    kwargs['year_on_year_east_90'][len(kwargs['index_east_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-float(
                    kwargs['year_on_year_east_90'][len(kwargs['index_east_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chain_east_90'][len(kwargs['index_east_90']) - 1]) > 0:
            tit = p.add_run('上涨' + str(
                float('%.2f' % (float(kwargs['chain_east_90'][len(kwargs['index_east_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('下降' + str(
                float('%.2f' % (-float(kwargs['chain_east_90'][len(kwargs['index_east_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        # 华中地区指数
        tit = p.add_run('华中地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['index_mid_90'][len(kwargs['index_mid_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['year_on_year_mid_90'][len(kwargs['index_mid_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (float(
                    kwargs['year_on_year_mid_90'][len(kwargs['index_mid_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-float(
                    kwargs['year_on_year_mid_90'][len(kwargs['index_mid_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chain_mid_90'][len(kwargs['index_mid_90']) - 1]) > 0:
            tit = p.add_run('上涨' + str(
                float('%.2f' % (float(kwargs['chain_mid_90'][len(kwargs['index_mid_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('下降' + str(
                float('%.2f' % (-float(kwargs['chain_mid_90'][len(kwargs['index_mid_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        # 华南地区指数
        tit = p.add_run('华南地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['index_south_90'][len(kwargs['index_south_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['year_on_year_south_90'][len(kwargs['index_south_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (float(
                    kwargs['year_on_year_south_90'][len(kwargs['index_south_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-float(
                    kwargs['year_on_year_south_90'][len(kwargs['index_south_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chain_south_90'][len(kwargs['index_south_90']) - 1]) > 0:
            tit = p.add_run('上涨' + str(
                float('%.2f' % (float(kwargs['chain_south_90'][len(kwargs['index_south_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('下降' + str(
                float('%.2f' % (-float(kwargs['chain_south_90'][len(kwargs['index_south_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        # 西南地区指数
        tit = p.add_run('西南地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(
            str(float('%.2f' % kwargs['index_southwest_90'][len(kwargs['index_southwest_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['year_on_year_southwest_90'][len(kwargs['index_southwest_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (float(
                    kwargs['year_on_year_southwest_90'][len(kwargs['index_southwest_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-float(
                    kwargs['year_on_year_southwest_90'][len(kwargs['index_southwest_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chain_southwest_90'][len(kwargs['index_southwest_90']) - 1]) > 0:
            tit = p.add_run('上涨' + str(
                float('%.2f' % (
                            float(kwargs['chain_southwest_90'][len(kwargs['index_southwest_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('下降' + str(
                float('%.2f' % (
                            -float(kwargs['chain_southwest_90'][len(kwargs['index_southwest_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        # 西北地区指数
        tit = p.add_run('西北地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(
            str(float('%.2f' % kwargs['index_northwest_90'][len(kwargs['index_northwest_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['year_on_year_northwest_90'][len(kwargs['index_northwest_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (float(
                    kwargs['year_on_year_northwest_90'][len(kwargs['index_northwest_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-float(
                    kwargs['year_on_year_northwest_90'][len(kwargs['index_northwest_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chain_northwest_90'][len(kwargs['index_northwest_90']) - 1]) > 0:
            tit = p.add_run('上涨' + str(
                float('%.2f' % (
                            float(kwargs['chain_northwest_90'][len(kwargs['index_northwest_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('下降' + str(
                float('%.2f' % (
                            -float(kwargs['chain_northwest_90'][len(kwargs['index_northwest_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        p = self.doc.add_paragraph()
        pic = p.add_run().add_picture(kwargs['index_by_sevenareas_image_url_90'])
        pic.height = Cm(6.17)
        pic.width = Cm(14.63)

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('附图3 七区域子市场走势图')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        # 四

        # kwargs['city_index_90_1']:一线城市“城房指数”的列表，按城市顺序排列，第一项为汇总的指数值
        # 'city_chain_90_1'以及'city_year_on_year_90_1'以此类推
        # kwargs['cityname_1']:一线城市名，第一项为“一线城市”，以此类推
        p = self.doc.add_paragraph()
        tit = p.add_run('四、各线城市子市场')
        tit.font.name = '黑体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        tit.font.size = Pt(14)

        # 一线城市
        p = self.doc.add_paragraph()
        tit = p.add_run('    各线城市子市场中，“一线城市”指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['city_index_90_1'][0])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['city_year_on_year_90_1'][0]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (
                        float(
                            kwargs['city_year_on_year_90_1'][0]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-float(
                    kwargs['city_year_on_year_90_1'][0]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['city_chain_90_1'][0]) > 0:
            tit = p.add_run('上涨' + str(
                float(
                    '%.2f' % (float(kwargs['city_chain_90_1'][0]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float('%.2f' % (-float(
                        kwargs['city_chain_90_1'][0]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        # 二线城市
        p = self.doc.add_paragraph()
        tit = p.add_run('“二线城市”指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['city_index_90_2'][0])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['city_year_on_year_90_2'][0]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (
                        float(
                            kwargs['city_year_on_year_90_2'][0]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-float(
                    kwargs['city_year_on_year_90_2'][0]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['city_chain_90_2'][0]) > 0:
            tit = p.add_run('上涨' + str(
                float(
                    '%.2f' % (float(kwargs['city_chain_90_2'][0]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float('%.2f' % (-float(
                        kwargs['city_chain_90_2'][0]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        # 三线城市
        p = self.doc.add_paragraph()
        tit = p.add_run('“三线城市”指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['city_index_90_3'][0])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['city_year_on_year_90_3'][0]) > 0:
            tit = p.add_run(
                '上涨' + str(float('%.2f' % (float(kwargs['city_year_on_year_90_3'][0]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float('%.2f' % (-float(kwargs['city_year_on_year_90_3'][0]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['city_chain_90_3'][0]) > 0:
            tit = p.add_run('上涨' + str(float('%.2f' % (float(kwargs['city_chain_90_3'][0]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('下降' + str(float('%.2f' % (-float(kwargs['city_chain_90_3'][0]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        # 四线城市
        p = self.doc.add_paragraph()
        tit = p.add_run('“四线城市”指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['city_index_90_4'][0])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['city_year_on_year_90_4'][0]) > 0:
            tit = p.add_run('上涨' + str(float('%.2f' % (float(kwargs['city_year_on_year_90_4'][0]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('下降' + str(float('%.2f' % (-float(kwargs['city_year_on_year_90_4'][0]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['city_chain_90_4'][0]) > 0:
            tit = p.add_run('上涨' + str(float('%.2f' % (float(kwargs['city_chain_90_4'][0]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run('下降' + str(float('%.2f' % (-float(kwargs['city_chain_90_4'][0]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        p = self.doc.add_paragraph()
        pic = p.add_run().add_picture(kwargs['index_by_line_image_url_90'])
        pic.height = Cm(6.17)
        pic.width = Cm(14.63)

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('附图4 各线城市子市场走势图')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        # 五
        p = self.doc.add_paragraph()
        tit = p.add_run('五、重点地区子市场')
        tit.font.name = '黑体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        tit.font.size = Pt(14)

        p = self.doc.add_paragraph()
        tit = p.add_run('    重点地区子市场中，珠三角地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['zhu_index_90'][len(kwargs['zhu_index_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['zhu_year_on_year_90'][len(kwargs['zhu_index_90']) - 1]) > 0:
            tit = p.add_run('上涨' + str(
                float(
                    '%.2f' % (float(kwargs['zhu_year_on_year_90'][len(kwargs['zhu_index_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float(
                    '%.2f' % (float(
                        -kwargs['zhu_year_on_year_90'][len(kwargs['zhu_index_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['zhu_chain_90'][len(kwargs['zhu_index_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(
                    float('%.2f' % (float(kwargs['zhu_chain_90'][len(kwargs['zhu_index_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float('%.2f' % (-float(kwargs['zhu_chain_90'][len(kwargs['zhu_index_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        tit = p.add_run('长三角地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['chang_index_90'][len(kwargs['chang_index_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['chang_year_on_year_90'][len(kwargs['chang_index_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(
                    float('%.2f' % (
                            float(kwargs['chang_year_on_year_90'][len(kwargs['chang_index_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float('%.2f' % (-float(
                        kwargs['chang_year_on_year_90'][len(kwargs['chang_index_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['chang_chain_90'][len(kwargs['chang_index_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(
                    float('%.2f' % (float(kwargs['chang_chain_90'][len(kwargs['chang_index_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float('%.2f' % (-float(kwargs['chang_chain_90'][len(kwargs['chang_index_90']) - 1]) * 100))) + '%；')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        tit = p.add_run('环渤海地区指数')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        tit = p.add_run(str(float('%.2f' % kwargs['bo_index_90'][len(kwargs['bo_index_90']) - 1])) + '点，同比')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)

        if float(kwargs['bo_year_on_year_90'][len(kwargs['bo_index_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(
                    float('%.2f' % (float(
                        kwargs['bo_year_on_year_90'][len(kwargs['bo_index_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(float(
                    '%.2f' % (-float(
                        kwargs['bo_year_on_year_90'][len(kwargs['bo_index_90']) - 1]) * 100))) + '%，环比')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        if float(kwargs['bo_chain_90'][len(kwargs['bo_index_90']) - 1]) > 0:
            tit = p.add_run(
                '上涨' + str(
                    float('%.2f' % (float(kwargs['bo_chain_90'][len(kwargs['bo_index_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)
        else:
            tit = p.add_run(
                '下降' + str(
                    float('%.2f' % (-float(kwargs['bo_chain_90'][len(kwargs['bo_index_90']) - 1]) * 100))) + '%。')
            tit.font.name = '仿宋_GB2312'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            tit.font.size = Pt(14)

        p = self.doc.add_paragraph()
        pic = p.add_run().add_picture(kwargs['index_by_focus_image_url_90_90'])
        pic.height = Cm(6.17)
        pic.width = Cm(14.63)

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('附图5 重点区域子市场走势图')
        tit.font.name = '仿宋_GB2312'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        tit.font.size = Pt(14)
        tit.font.bold = True

        # 附注

        p = self.doc.add_paragraph()
        tit = p.add_run('附注')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(10)
        tit.font.bold = True

        p = self.doc.add_paragraph()
        tit = p.add_run('（1）东部地区包括北京、天津、石家庄、沈阳、大连、上海、南京、无锡、苏州、杭州、宁波、温州、福州、厦门、济南、青岛、'
                        '广州、深圳、海口、三亚、秦皇岛、唐山、本溪、丹东、锦州、烟台、淄博、泰安、潍坊、扬州、徐州、常州、金华、绍兴、泉州、佛'
                        '山、汕头、中山、惠州、江门、湛江、韶关等42个城市。中部地区包括太原、长春、哈尔滨、合肥、南昌、郑州、武汉、长沙、大同、'
                        '齐齐哈尔、牡丹江、大庆、吉林、宣城、蚌埠、安庆、芜湖、赣州、九江、南阳、洛阳、平顶山、宜昌、襄阳、株洲、岳阳、常德等27'
                        '个城市。西部地区包括呼和浩特、南宁、北海、重庆、成都、贵阳、昆明、西安、兰州、西宁、银川、乌鲁木齐、包头、桂林、大理、'
                        '遵义、绵阳、南充、泸州、宝鸡、柳州等21个城市。')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(10)

        p = self.doc.add_paragraph()
        tit = p.add_run('（2）东北地区包括沈阳、大连、长春、哈尔滨、齐齐哈尔、牡丹江、大庆、吉林、本溪、丹东、锦州等11个城市。华北地区包括北京、'
                        '天津、石家庄、太原、呼和浩特、秦皇岛、唐山、大同、包头等9个城市。华东地区包括上海、南京、无锡、苏州、杭州、宁波、温州、'
                        '合肥、济南、青岛、烟台、淄博、泰安、潍坊、扬州、徐州、常州、宣城、蚌埠、安庆、芜湖、金华、绍兴等23个城市。华中地区包括'
                        '南昌、郑州、武汉、长沙、赣州、九江、南阳、洛阳、平顶山、宜昌、襄阳、株洲、岳阳、常德等14个城市。华南地区包括福州、厦门'
                        '、广州、深圳、南宁、北海、海口、三亚、泉州、佛山、汕头、中山、惠州、江门、湛江、韶关、桂林、柳州等18个城市。西南地区包'
                        '括重庆、成都、贵阳、昆明、大理、遵义、绵阳、南充、泸州等9个城市。西北地区包括西安、兰州、西宁、银川、乌鲁木齐、宝鸡等6个城市。')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(10)

        p = self.doc.add_paragraph()
        tit = p.add_run('（3）一线城市包括北京、上海、广州、深圳等4个城市。二线城市包括天津、石家庄、太原、呼和浩特、沈阳、大连、长春、哈尔滨、'
                        '南京、无锡、苏州、杭州、宁波、温州、合肥、福州、厦门、南昌、济南、青岛、郑州、武汉、长沙、南宁、北海、海口、三亚、重庆'
                        '、成都、贵阳、昆明、西安、兰州、西宁、银川、乌鲁木齐等36个城市。三线城市包括秦皇岛、唐山、包头、齐齐哈尔、大庆、吉林、'
                        '锦州、烟台、淄博、潍坊、扬州、徐州、常州、蚌埠、安庆、芜湖、金华、绍兴、泉州、赣州、九江、洛阳、宜昌、襄阳、株洲、岳阳、'
                        '常德、佛山、中山、惠州、江门、湛江、韶关、桂林、绵阳、南充、泸州等37个城市。四线城市包括大同、牡丹江、本溪、丹东、泰安、'
                        '宣城、南阳、平顶山、汕头、大理、遵义、宝鸡、柳州等13个城市。')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(10)

        p = self.doc.add_paragraph()
        tit = p.add_run('（4）珠三角区域包括广州、深圳、佛山、中山、惠州、江门等6个城市。长三角区域包括上海、南京、无锡、苏州、杭州、宁波、温州'
                        '、扬州、常州、金华、绍兴等11个城市。环渤海区域包括北京、天津、石家庄、沈阳、大连、青岛、秦皇岛、唐山、丹东、锦州、烟台'
                        '、潍坊等12个城市。')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(10)

    def EndReport(self):
        self.doc.save(self.url)


class GenWordPicture90:

    def __init__(self, url):
        self.url = url
        self.doc = Document()

    def genpicdoc(self, year: int, month: int, citylist: list, **kwargs):
        # kwargs['city_index_url_90']:list city index
        # kwargs['total_index_url_90']:str
        # kwargs['total_year_on_year_url_90']:str
        # kwargs['total_chain_url_90']:str
        # kwargs['index_block_90']:str
        # kwargs['index_area_90']:str

        for i in range(90):
            p = self.doc.add_paragraph()
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            pic = p.add_run().add_picture(kwargs['city_index_url_90'][i])
            pic.height = Cm(8.11)
            pic.width = Cm(14.01)
            p = self.doc.add_paragraph()
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            tit = p.add_run(
                '图' + str(i + 1) + '  ' + citylist[i] + '市“城房指数”（2009.1~' + str(year) + '.' + str(month) + '）')
            tit.font.name = '宋体'
            tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            tit.font.size = Pt(10.5)

        # 图91

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pic = p.add_run().add_picture(kwargs['total_index_url_90'])
        pic.height = Cm(7.98)
        pic.width = Cm(19.03)
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('图91  90城市“城房指数”（2009.1~' + str(year) + '.' + str(month) + '）')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(10.5)

        # 图92

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pic = p.add_run().add_picture(kwargs['total_year_on_year_url_90'])
        pic.height = Cm(4.83)
        pic.width = Cm(17.06)
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('图92  90城市“城房指数”同比变化（2010.1~' + str(year) + '.' + str(month) + '）')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(10.5)

        # 图93

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pic = p.add_run().add_picture(kwargs['total_chain_url_90'])
        pic.height = Cm(4.83)
        pic.width = Cm(17.06)
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('图93  90城市“城房指数”环比变化（2009.2~' + str(year) + '.' + str(month) + '）')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(10.5)

        # 图94

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pic = p.add_run().add_picture(kwargs['index_block_90'])
        pic.height = Cm(8.17)
        pic.width = Cm(18.99)
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('图94  90城市东中西区域“城房指数”（2009.1~' + str(year) + '.' + str(month) + '）')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(10.5)

        # 图95

        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pic = p.add_run().add_picture(kwargs['index_area_90'])
        pic.height = Cm(8.17)
        pic.width = Cm(18.99)
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tit = p.add_run('图95  90城市各面积子市场“城房指数”（2009.1~' + str(year) + '.' + str(month) + '）')
        tit.font.name = '宋体'
        tit.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tit.font.size = Pt(10.5)

    def EndReport(self):
        self.doc.save(self.url)
